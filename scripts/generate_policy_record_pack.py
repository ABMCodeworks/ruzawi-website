from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from scripts.generate_compliance_pack import (
        ROOT, LOGO, GREEN, BLUE, GOLD, INK, MID, PALE_GREEN, WHITE,
        setup_styles, set_run_font, set_table_geometry, shade,
        add_page_field, add_hyperlink, add_h, add_p, add_bullets,
        add_numbers, add_callout, add_table,
    )
except ModuleNotFoundError:
    from generate_compliance_pack import (
        ROOT, LOGO, GREEN, BLUE, GOLD, INK, MID, PALE_GREEN, WHITE,
        setup_styles, set_run_font, set_table_geometry, shade,
        add_page_field, add_hyperlink, add_h, add_p, add_bullets,
        add_numbers, add_callout, add_table,
    )


OUT = ROOT / "docs" / "potraz-policy-record-pack"
ACT_URL = "https://www.potraz.gov.zw/wp-content/uploads/2026/02/ACT-CDPA.pdf"
REG_URL = "https://www.potraz.gov.zw/wp-content/uploads/2025/02/sI-155-of-2024-Cyber-and-Data-Protection-Normal_240913_1250178.pdf"
POTRAZ_URL = "https://www.potraz.gov.zw/"
CHILD_URL = "https://www.veritaszim.net/sites/veritas_d/files/02%20Processing%20of%20Children%27s%20Personal%20Information.pdf"


def new_policy(title, subtitle, policy_id, owner):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.82)
    sec.bottom_margin = Inches(0.76)
    sec.left_margin = Inches(0.86)
    sec.right_margin = Inches(0.86)
    sec.header_distance = Inches(0.32)
    sec.footer_distance = Inches(0.35)
    setup_styles(doc)

    header = sec.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"RUZAWI SCHOOL  |  {policy_id}")
    set_run_font(r, 8, bold=True, color=BLUE)

    footer = sec.footer
    p = footer.paragraphs[0]
    r = p.add_run("Controlled policy record - approve, communicate and retain with Board minutes")
    set_run_font(r, 8, color=MID)
    add_page_field(footer.add_paragraph())

    if LOGO.exists():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        image = p.add_run().add_picture(str(LOGO), width=Inches(0.72))
        image._inline.docPr.set("descr", "Ruzawi School crest")
        image._inline.docPr.set("title", "Ruzawi School")

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    r = kicker.add_run("POTRAZ COMPLIANCE POLICY RECORD")
    set_run_font(r, 9, bold=True, color=GOLD)
    doc.add_paragraph(title, style="Title")
    doc.add_paragraph(subtitle, style="Subtitle")
    return doc


def _policy_bullet_num_id(doc):
    cached = getattr(doc, "_policy_bullet_num_id", None)
    if cached is not None:
        return cached

    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(value)
        for value in numbering.xpath("./w:abstractNum/@w:abstractNumId")
    ]
    num_ids = [int(value) for value in numbering.xpath("./w:num/@w:numId")]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "–")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "420")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "420")
    indent.set(qn("w:hanging"), "240")
    p_pr.append(indent)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    doc._policy_bullet_num_id = num_id
    return num_id


def add_bullets(doc, items):
    num_id = _policy_bullet_num_id(doc)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_pr.append(ilvl)
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(num_id_el)
        p_pr.append(num_pr)
        r = p.add_run(item)
        set_run_font(r, 10.5)


def add_common_end(doc, evidence, children_source=False):
    add_h(doc, "Records retained as evidence of implementation", 1)
    add_bullets(doc, evidence)
    add_h(doc, "Monitoring, non-compliance and review", 1)
    add_p(doc, "The policy owner shall monitor implementation, report material exceptions to the Data Protection Officer and Head, and ensure corrective action is assigned and completed. Deliberate or negligent non-compliance may result in access suspension, disciplinary action, contractual remedies, notification to POTRAZ or other legal action. The Data Protection Officer shall review this policy at least annually and whenever the governing law, POTRAZ guidance, School processing, technology, suppliers or risk materially changes.")
    legal_heading = add_h(doc, "Legal basis and related records", 1)
    legal_heading.paragraph_format.keep_with_next = True
    sources = [
        ("Cyber and Data Protection Act [Chapter 12:07]", ACT_URL),
        ("S.I. 155 of 2024", REG_URL),
        ("POTRAZ Data Protection Authority", POTRAZ_URL),
    ]
    if children_source:
        sources.append(("Implementation Guideline on Processing Children's Personal Information (CDPG 2 of 2024)", CHILD_URL))
    for name, url in sources:
        p = doc.add_paragraph(style="Small Text")
        p.paragraph_format.keep_with_next = True
        add_hyperlink(p, name, url)
    p = doc.add_paragraph(style="Small Text")
    r = p.add_run("Related internal framework: Ruzawi public Privacy Policy and Cookie Policy; the other policies in this controlled POTRAZ policy record pack; the School's current DP1/DP2/DP3 filings, ROPA, registers, contracts and implementation evidence.")
    set_run_font(r, 8.5, color=MID)


def save(doc, filename):
    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / filename
    doc.core_properties.title = filename.replace(".docx", "")
    doc.core_properties.subject = "Ruzawi School POTRAZ compliance policy record"
    doc.core_properties.author = "Ruzawi School"
    doc.core_properties.keywords = "POTRAZ, data protection, policy, Ruzawi School"
    doc.save(path)
    return path


def governance_policy():
    doc = new_policy("Data Protection Governance and Accountability Policy", "The School's overarching framework for responsibility, oversight and demonstrable compliance", "POL-DP-01", "Data Protection Officer")
    add_callout(doc, "Policy position", "Ruzawi School accepts accountability for all personal information processed under its control and will maintain the internal mechanisms, records and management oversight needed to demonstrate compliance to data subjects and POTRAZ.")
    add_h(doc, "1. Purpose and scope", 1)
    add_p(doc, "This policy establishes the School's governance framework for personal information relating to pupils, parents and guardians, staff, applicants, alumni, visitors, governors, volunteers, suppliers and other individuals. It applies to all paper and electronic processing by the School and by processors acting on its instructions.")
    add_h(doc, "2. Governance principles", 1)
    add_bullets(doc, [
        "Processing must be lawful, fair, transparent, necessary, proportionate, accurate, secure and limited to defined purposes and retention periods.",
        "Children's best interests and privacy must be considered from the outset and protected by design and by default.",
        "The School must be able to prove implementation through current policies, registrations, ROPA entries, approvals, contracts, logs, training and review records.",
        "No department or employee may introduce a new processing purpose, system, supplier, surveillance technology, international transfer or automated decision without the required privacy review.",
    ])
    add_h(doc, "3. Data Protection Officer", 1)
    add_p(doc, "The Head and Board shall appoint a suitably qualified Data Protection Officer, notify POTRAZ using the prescribed process and provide adequate authority, independence, resources and access to senior management. The DPO shall monitor compliance, manage internal privacy activities, raise awareness, train staff, conduct or coordinate audits, handle requests from data subjects and POTRAZ, advise on DPIAs and act as the contact point for the Authority and data subjects.")
    add_h(doc, "4. Leadership and departmental accountability", 1)
    add_table(doc, ["Role", "Accountability"], [
        ("Board and Head", "Approve the framework, appoint and resource the DPO, review material risks/incidents and require remediation."),
        ("DPO", "Monitor, advise, train, audit, maintain governance records, oversee requests/DPIAs/incidents and liaise with POTRAZ."),
        ("Information owners", "Maintain accurate processing records, lawful basis, access, notices, retention, risks and evidence for their functions."),
        ("IT and Security", "Implement and test technical, physical and continuity controls; support incident response and system evidence."),
        ("All personnel", "Use information only as authorised, follow policies, complete training and immediately report incidents and rights requests."),
    ], [2200, 7160], font=8.6)
    add_h(doc, "5. Required compliance framework", 1)
    add_p(doc, "The DPO shall maintain a controlled policy set, current data-controller licence and regulatory filings, a Record of Processing Activities, consent and rights records, breach records, processor and transfer records, retention evidence, DPIAs, risk assessments, training records, audit findings and evidence of Board oversight. Records must be available for inspection within a reasonable period and protected according to sensitivity.")
    add_h(doc, "6. Assurance and reporting", 1)
    add_p(doc, "The DPO shall conduct quarterly compliance reviews and an annual internal assessment against the current POTRAZ framework. A written annual report shall be presented to the Head and Board covering licensing, requests, incidents, DPIAs, transfers, processors, training, audits, risks and overdue remediation. Material breaches, unlawful processing or critical residual risks must be escalated immediately rather than waiting for the routine report.")
    add_common_end(doc, [
        "Current DP1 application or renewal, data-controller licence, fee receipts and renewal calendar.",
        "DPO appointment instrument, qualifications or training record, DP2 notification and POTRAZ acknowledgement.",
        "Approved policy set, version history, Board/SMT approval minutes and annual DPO reports.",
        "Current organisation chart, role descriptions, audit reports, remediation records and evidence of resourcing.",
    ])
    return save(doc, "01-Data-Protection-Governance-and-Accountability-Policy.docx")


def lawful_processing_policy():
    doc = new_policy("Lawful Processing, Consent and Data Minimisation Policy", "Rules for lawful basis, transparency, consent, purpose limitation and minimum collection", "POL-DP-02", "Data Protection Officer and Information Owners")
    add_callout(doc, "Policy position", "Personal information may be processed only for a specific recorded purpose, on a documented lawful basis, using the minimum information necessary. Consent will be used only when it is genuine, specific and capable of withdrawal.")
    add_h(doc, "1. Lawful basis", 1)
    add_p(doc, "Before processing begins, the information owner must record the purpose, necessity, data subjects, data categories, source, recipients, systems, transfers, retention, safeguards and lawful basis in the School's ROPA. Depending on the activity, the basis may include valid consent, compliance with law, protection of vital interests, performance of a task in the public interest, legitimate interests not overridden by the individual's rights, or another basis expressly permitted by law. Sensitive information requires written consent unless a documented statutory exception applies.")
    add_h(doc, "2. Consent standard", 1)
    add_bullets(doc, [
        "Consent must be informed, specific to a defined purpose, freely given, recorded and expressed through a clear affirmative act.",
        "Optional consent must not be bundled with enrolment, employment or another service when the information is not necessary for that relationship.",
        "The School must preserve the wording and version shown, the consenting person and authority, the purpose, date, method and withdrawal history.",
        "Consent may be withdrawn at any time, without explanation and free of charge. Future consent-based processing must stop promptly unless another lawful basis is separately documented.",
        "Changes to purpose, data, recipients, technology, transfer or risk require renewed assessment and, where necessary, refreshed consent and notice.",
    ])
    add_h(doc, "3. Transparency", 1)
    add_p(doc, "At or before collection, the School shall provide a clear notice identifying Ruzawi as controller, its address and DPO contact, the purposes, required and optional fields, consequences of not providing required information, recipients, transfers, retention approach and available rights. Where data is obtained indirectly, the individual must be informed unless a lawful exception applies. Notices must be current, accessible and consistent with actual practice.")
    add_h(doc, "4. Purpose limitation and minimisation", 1)
    add_p(doc, "Information may not be repurposed merely because it is available. Any new use requires written compatibility and lawful-basis review by the information owner and DPO. Forms and systems must distinguish mandatory from optional fields, avoid unrestricted free-text collection, limit uploads, and prevent staff from retaining duplicate local copies. Departments shall review material forms and system fields at least annually and remove fields that are not demonstrably necessary.")
    add_h(doc, "5. Accuracy and correction", 1)
    add_p(doc, "Material information must be kept accurate and updated where necessary. Individuals shall be given appropriate opportunities to correct relevant data. When an error is substantiated, the School must correct or delete it without delay across relevant systems and notify processors or recipients where appropriate.")
    add_h(doc, "6. Marketing and public information", 1)
    add_p(doc, "Application, pupil, parent or employment information may not be used for unrelated direct marketing without valid consent or another documented authority. A person's own publication of information does not by itself authorise the School to collect or reuse it. Image, publication, alumni and fundraising uses must be recorded by purpose and choice, with suppression applied when consent is withdrawn.")
    add_common_end(doc, [
        "Current ROPA and lawful-basis decisions for each processing purpose.",
        "Approved privacy notices, consent wording/version records and withdrawal evidence.",
        "Annual form-field and system-minimisation reviews, purpose-change approvals and correction records.",
        "Samples demonstrating that optional choices do not affect admission, employment or access to core services.",
    ], children_source=True)
    return save(doc, "02-Lawful-Processing-Consent-and-Data-Minimisation-Policy.docx")


def rights_policy():
    doc = new_policy("Data Subject Rights Policy", "School rules for access, objection, correction, deletion, consent withdrawal and complaints", "POL-DP-03", "Data Protection Officer")
    add_callout(doc, "Policy position", "Ruzawi will recognise privacy requests in any reasonable form, verify identity and authority proportionately, search all relevant records, protect third parties and provide a clear, secure and timely response.")
    add_h(doc, "1. Rights recognised", 1)
    add_p(doc, "Subject to the Act and applicable limitations, individuals have rights to be informed about use, access personal information held by the School, object to all or part of processing, correct false, misleading or outdated information, delete false or misleading information, withdraw consent without charge and object without charge to direct marketing. A parent or legal guardian may exercise a child's rights.")
    add_h(doc, "2. Intake and responsibility", 1)
    add_p(doc, "The DPO owns all rights requests. privacy@ruzawi.com must be monitored, protected by multi-factor authentication and covered by a trained alternate. A request received by any member of staff must be forwarded immediately; the requester may not be required to use a particular form or legal terminology.")
    add_h(doc, "3. Verification and authority", 1)
    add_p(doc, "Verification must be proportionate to the sensitivity and disclosure risk. The School shall not request or retain excessive identity documents. Where a request is made for a child, the School must verify parental or legal-guardian authority and consider safeguarding restrictions or conflicting instructions before disclosure. Other representatives must provide reliable authority.")
    add_h(doc, "4. Search, decision and response", 1)
    add_numbers(doc, [
        "Record the request, receipt date, scope, verification status, responsible owners and internal due date.",
        "Search all relevant electronic, paper, email, cloud, CCTV, archived and processor-held locations using reliable identifiers and date ranges.",
        "Review the information for accuracy, legal holds, safeguarding concerns, confidentiality and the rights of other people.",
        "Complete corrections, deletions, objections or consent withdrawal across all relevant systems and processors where the request is valid.",
        "Provide a clear decision and information securely, explaining any lawful refusal or limitation and the route to complain to the DPO and POTRAZ.",
    ])
    add_h(doc, "5. Timeliness", 1)
    add_p(doc, "The School adopts 30 calendar days from receipt as its internal control period for completion, consistent with the compliance checklist supplied to the School. Verification and clarification must be pursued promptly and may not be used to create avoidable delay. Any request at risk of delay must be escalated to the DPO and Head with a recorded recovery plan.")
    add_h(doc, "6. Secure disclosure and recordkeeping", 1)
    add_p(doc, "Responses must use a channel appropriate to sensitivity. The destination must be verified; passwords or keys must be exchanged separately when encryption is used. The School shall retain a restricted record of the request, searches, decision, actions, response and delivery evidence for accountability and dispute handling, subject to the approved retention schedule.")
    add_common_end(doc, [
        "Rights-request register showing receipt, internal due date, scope, action, response and closure.",
        "Identity and guardian-authority verification outcomes without unnecessary copies.",
        "Search instructions and confirmations from relevant information owners and processors.",
        "Redacted examples of timely completed requests, secure delivery and corrections or deletions implemented.",
    ], children_source=True)
    return save(doc, "03-Data-Subject-Rights-Policy.docx")


def children_policy():
    doc = new_policy("Children's Personal Information Protection Policy", "Best interests, guardian authority, age-appropriate transparency and high-risk controls", "POL-DP-04", "Data Protection Officer and Safeguarding Lead")
    add_callout(doc, "Policy position", "Because pupils are children, Ruzawi will place their privacy and best interests at the centre of every processing decision and apply stronger controls than it ordinarily applies to adult data.")
    add_h(doc, "1. Scope and governing standard", 1)
    add_p(doc, "This policy applies to information about any person under 18, including admissions, academic, pastoral, boarding, health, safeguarding, financial, photographic, biometric, CCTV, device, communications and activity information. Children's rights under the Act may be exercised by their parents or legal guardians.")
    add_h(doc, "2. Consent and guardian verification", 1)
    add_p(doc, "Where consent is required, it must be written, specific and obtained from a parent or legal guardian before processing. The School must make reasonable efforts to verify the adult's authority through an appropriate birth certificate, adoption, custody or guardianship record or other reliable evidence. The School should record the evidence sighted and verification result, retaining a full copy only when necessary and lawful.")
    add_h(doc, "3. Age-appropriate information", 1)
    add_p(doc, "Privacy notices, consent requests and explanations must be understandable to the parent or guardian and appropriate to the pupil's age and maturity. Pupils must be told in simple language what is collected, why, who may receive it, how it is protected and how to raise a question or concern. Staff must discourage unnecessary over-sharing and must not treat a child's public social-media activity as permission for unrelated processing.")
    add_h(doc, "4. Best interests and privacy by design", 1)
    add_bullets(doc, [
        "Collect and expose the minimum data; use privacy-protective defaults and restrict search, sharing and publication.",
        "Screen every new pupil-facing system or supplier before procurement and complete a DPIA where the activity may create high risk.",
        "Conduct children's-data-specific supplier due diligence and ensure written processor obligations and deletion controls.",
        "Separate education, safeguarding, health and disciplinary access so that sensitive information is available only to staff with a genuine need.",
        "Provide meaningful human consideration of the child's circumstances and ability to challenge material decisions.",
    ])
    add_h(doc, "5. High-risk processing and POTRAZ engagement", 1)
    add_p(doc, "New surveillance or tracking technology, biometric use, cross-border transfer or processing capable of causing discrimination, identity theft, financial, reputational or serious privacy harm must be classified as high risk. The School shall complete a DPIA and obtain prior POTRAZ authorisation or a documented exemption where required before the activity begins. Where a child's data is processed without consent on another lawful basis, the DPO shall document the basis and notify POTRAZ in writing where required by CDPG 2 of 2024.")
    add_h(doc, "6. Automated decisions", 1)
    add_p(doc, "A child shall not be subject to a decision based solely on automated processing that produces legal or similarly significant effects unless expressly permitted by law and approved by the DPO. Any system affecting admission, placement, learning support, discipline, safeguarding or access must provide meaningful human review, the ability to challenge and a clear explanation of the decision.")
    add_h(doc, "7. Age 18 and leaver review", 1)
    add_p(doc, "When a pupil turns 18 or otherwise ceases to be a child, the relevant information owner shall review information and permissions: delete what is no longer necessary, retain only under a recorded lawful basis and period, and obtain consent directly from the adult where future optional processing is desired. The decision and completion across relevant systems and processors must be recorded.")
    add_common_end(doc, [
        "Guardian-authority verification records and versioned parental consent evidence.",
        "Age-appropriate pupil notice and evidence of communication or explanation.",
        "DPIAs, POTRAZ authorisations or exemptions and non-consent notifications for applicable activities.",
        "Supplier reviews, automated-decision human-review evidence and age-18/leaver review records.",
    ], children_source=True)
    return save(doc, "04-Childrens-Personal-Information-Protection-Policy.docx")


def security_policy():
    doc = new_policy("Information Security and Access Control Policy", "Minimum technical, organisational and physical safeguards for personal information", "POL-DP-05", "IT Security Lead and Data Protection Officer")
    add_callout(doc, "Policy position", "Ruzawi will protect confidentiality, integrity, availability and recoverability using controls proportionate to the sensitivity of the information and the harm that a failure could cause, especially to children.")
    add_h(doc, "1. Scope", 1)
    add_p(doc, "This policy applies to School networks, email, cloud services, management systems, websites, servers, devices, removable media, paper records, CCTV, backups and processors. It applies whether information is used on campus, remotely or by a supplier.")
    add_h(doc, "2. Identity and access", 1)
    add_bullets(doc, [
        "Every user must have a unique account; shared credentials are prohibited except for formally controlled service accounts.",
        "Multi-factor authentication must be enabled for email, administration, remote access, School management systems and other sensitive services wherever available.",
        "Access must follow least privilege, require owner approval and be removed promptly when a person leaves or changes role.",
        "Sensitive-system access shall be reviewed at least termly and privileged access at least quarterly, with removals and exceptions recorded.",
    ])
    add_h(doc, "3. Device, network and cloud security", 1)
    add_p(doc, "School-managed devices must use supported software, secure configuration, timely patches, screen lock, malware protection and encryption where proportionate. Administration must be restricted and logged. Remote access and cloud services require approved security, supplier and transfer review. Personal devices and consumer storage or messaging may not be used for restricted or confidential School information without written approval and compensating controls.")
    add_h(doc, "4. Information handling", 1)
    add_p(doc, "Restricted information, including health, safeguarding, credentials, biometrics and identity documents, must be limited to named roles, securely transmitted and protected from uncontrolled download, printing or forwarding. Staff must verify recipients, use approved sharing methods, collect printouts promptly and keep records secure from pupils, visitors and unauthorised staff.")
    add_h(doc, "5. Logging, backup and continuity", 1)
    add_p(doc, "Relevant authentication, administrative and data-access events must be logged and protected from alteration. Alerts must be reviewed and investigated. Critical information must be backed up according to recovery need, protected from production compromise and tested through documented restoration exercises. Security controls must be tested and improved rather than assumed to be effective.")
    add_h(doc, "6. Physical security", 1)
    add_p(doc, "Records rooms, server/network areas, cabinets and keys must be restricted to authorised personnel. Visitors must be supervised or logged as appropriate. Confidential records may not be left unattended in classrooms, vehicles or shared areas. Physical transfer and storage must protect against loss, theft, fire, water and unauthorised access in proportion to risk.")
    add_h(doc, "7. Security exceptions and reporting", 1)
    add_p(doc, "Any exception requires a recorded reason, scope, risk, compensating control, owner, approval and expiry date. Indefinite exceptions are prohibited. Lost devices, suspicious messages, misdirected email, unauthorised access, malware, missing paper records and other suspected incidents must be reported immediately under the Data Breach Response and Notification Policy.")
    add_common_end(doc, [
        "Access-control matrices, approvals, MFA reports and joiner/mover/leaver records.",
        "Patch, endpoint, encryption, vulnerability, security-configuration and alert evidence.",
        "Backup success and restoration tests, incident records and control-improvement actions.",
        "Records-room, key, visitor and privileged-access review records; approved security exceptions.",
    ])
    return save(doc, "05-Information-Security-and-Access-Control-Policy.docx")


def cctv_policy():
    doc = new_policy("CCTV Surveillance and Retention Policy", "Purpose, access, disclosure, high-risk children's processing and a 30-day routine retention rule", "POL-DP-06", "Head of Security and Data Protection Officer")
    add_callout(doc, "Policy position", "CCTV will be used only for defined safety, safeguarding and security purposes, with visible notice, restricted access, a documented DPIA and routine footage automatically overwritten after 30 days.")
    add_h(doc, "1. Permitted purpose and placement", 1)
    add_p(doc, "Cameras may be installed only where necessary and proportionate for pupil, staff, visitor or property safety, safeguarding or security. Cameras must not be positioned in toilets, changing areas or other locations where privacy expectations make surveillance disproportionate. Each camera and viewing purpose must be recorded in the CCTV inventory and DPIA.")
    add_h(doc, "2. Transparency and children's safeguards", 1)
    add_p(doc, "Clear signage shall identify CCTV use and the School as controller and provide a privacy contact route. Because the footage usually concerns children, the DPO shall assess best interests, field of view, recording hours, masking, access, transfers and alternatives. Required POTRAZ authorisation or exemption for high-risk children's processing must be obtained before a new or materially changed system is used.")
    add_h(doc, "3. Recording and prohibited functions", 1)
    add_p(doc, "The system shall record only what is necessary for the approved purpose. Audio recording, facial recognition, biometric identification, emotion analysis or automated behavioural profiling is prohibited unless separately justified, legally reviewed, subject to DPIA, approved by the Head and Board and authorised by POTRAZ where required.")
    add_h(doc, "4. Access, viewing and export", 1)
    add_p(doc, "Live and recorded access is limited to named authorised roles using unique accounts and secure credentials. Viewing, search, export and disclosure must be logged. Exports require a case reference, minimum relevant time segment, approved encrypted storage and a defined review/deletion trigger. Footage may not be copied to personal devices, consumer messaging or uncontrolled removable media.")
    add_h(doc, "5. Retention schedule", 1)
    add_table(doc, ["CCTV record", "Retention rule", "Required control"], [
        ("Routine footage", "30 days from recording", "Automatic overwrite or deletion; Security verifies and records the configured period monthly."),
        ("Incident or safeguarding export", "Until the investigation, proceeding, claim, appeal and legal hold are closed, followed by prompt deletion unless a longer lawful period is recorded", "Restricted case file, access/disclosure log and review at least every 90 days."),
        ("Footage preserved for POTRAZ, law enforcement, insurer, court or rights request", "Only while the authorised request, legal duty or dispute remains active", "DPO or legal review; preserve original integrity; disclose the minimum necessary."),
        ("CCTV access and export logs", "Six years proposed from entry or case closure, subject to final legal approval", "Restricted compliance record and secure deletion at expiry."),
    ], [2200, 3700, 3460], font=8.0)
    add_h(doc, "6. Holds and disclosures", 1)
    add_p(doc, "A hold must identify the exact camera/time range, reason, owner, review date and release trigger. Blanket or indefinite preservation is prohibited. Disclosure must be lawful, proportionate and recorded, with third-party privacy and safeguarding risks considered. Requests from data subjects are handled under the Data Subject Rights Policy.")
    add_h(doc, "7. Review and system assurance", 1)
    add_p(doc, "Security and the DPO shall review camera purpose, field of view, signage, access, retention configuration, open holds, exports, disclosures and security at least annually and after any incident or material change. Unnecessary cameras, access or recorded functions shall be removed or disabled.")
    add_common_end(doc, [
        "Approved CCTV inventory, camera map, privacy notice/signage and completed DPIA.",
        "POTRAZ authorisation or documented exemption where required for children's high-risk processing.",
        "Named access list, viewing/export/disclosure logs and monthly 30-day overwrite verification.",
        "Incident holds, 90-day reviews, rights/law-enforcement decisions and secure deletion evidence.",
    ], children_source=True)
    return save(doc, "06-CCTV-Surveillance-and-Retention-Policy.docx")


def breach_policy():
    doc = new_policy("Data Breach Response and Notification Policy", "Immediate reporting, containment, POTRAZ notification, affected-person communication and investigation", "POL-DP-07", "Data Protection Officer and IT Security Lead")
    add_callout(doc, "Policy position", "Every suspected personal-data incident must be reported immediately. Incomplete facts are not a reason to delay escalation, containment or the statutory notification timetable.", "warning")
    add_h(doc, "1. Events covered", 1)
    add_p(doc, "A personal-data incident includes accidental or unlawful loss, destruction, alteration, unauthorised access or disclosure, misdirected communication, missing paper file or device, malware, compromised account, processor incident, ransomware or material unavailability affecting personal information.")
    add_h(doc, "2. Immediate duties", 1)
    add_numbers(doc, [
        "The person discovering the event reports it immediately to the DPO and IT Security using the approved emergency route.",
        "IT contains the event safely, preserves logs and evidence, prevents further loss and avoids actions that would compromise investigation.",
        "The DPO records discovery and awareness times, appoints the incident team and determines whether personal information is affected.",
        "The team identifies affected systems, data, people, children or vulnerable individuals, recipients, protection status, likely consequences and ongoing risk.",
    ])
    add_h(doc, "3. Notification duties", 1)
    add_bullets(doc, [
        "The DPO shall report a personal data breach to POTRAZ within 24 hours of the School becoming aware, using Form DP3 and preserving proof of submission.",
        "Where the breach is likely to result in high risk to individuals' rights and freedoms, affected people shall be informed within 72 hours in clear, practical and age-appropriate language.",
        "The School shall respond to POTRAZ requests concerning the breach within 14 days.",
        "The investigation shall be concluded and a final report submitted within 21 days of the initial notification.",
    ])
    add_h(doc, "4. Risk and communication", 1)
    add_p(doc, "Risk assessment shall consider sensitivity, identifiability, encryption, volume, vulnerability, children, health, safeguarding, credentials, financial information, likely misuse, physical safety, distress, identity theft, discrimination, duration and reversibility. Communications must explain what happened, likely consequences, action taken, practical protective steps, DPO contact and how updates will be provided, without exposing additional personal information.")
    add_h(doc, "5. Investigation and recovery", 1)
    add_p(doc, "The incident team shall establish chronology and root cause, identify accounts/systems and processor involvement, eradicate malicious access, rotate credentials, correct configuration, restore clean and available data, test controls and preserve a controlled evidence file. Corrective actions must update relevant risks, DPIAs, contracts, training, retention, access and policies.")
    add_h(doc, "6. Cooperation and confidentiality", 1)
    add_p(doc, "Processors must notify Ruzawi without undue delay and provide facts, logs and assistance sufficient for the School's deadlines. Staff must not conceal, speculate publicly or contact affected people independently. Communications with POTRAZ, affected people, insurers, law enforcement and the media shall be coordinated by authorised leaders, the DPO and legal adviser as appropriate.")
    add_h(doc, "7. Exercises", 1)
    add_p(doc, "The DPO and IT Security shall conduct at least one documented breach tabletop exercise annually, including a scenario involving children's sensitive information. Lessons and remediation shall be reported to the Head and Board and tracked to closure.")
    add_common_end(doc, [
        "Breach register with discovery/awareness time, data, risk, decisions, containment and closure.",
        "Submitted DP3, delivery acknowledgement, affected-person notices and all POTRAZ correspondence.",
        "Controlled investigation file, evidence chain, root-cause and 21-day final report.",
        "Tabletop exercise report, remediation plan, control changes and Board/SMT incident oversight.",
    ], children_source=True)
    return save(doc, "07-Data-Breach-Response-and-Notification-Policy.docx")


def retention_policy():
    doc = new_policy("Records Retention and Secure Destruction Policy", "Approved retention principles, category schedule, legal holds and document shredding", "POL-DP-08", "Data Protection Officer and Records Owners")
    add_callout(doc, "Policy position", "Personal information will not be kept indefinitely. Every category must have a defined purpose, trigger and retention period, followed by secure deletion, destruction or irreversible anonymisation unless an authorised hold applies.")
    add_h(doc, "1. Retention rules", 1)
    add_bullets(doc, [
        "The shortest applicable period shall be used after considering legal, education, safeguarding, employment, tax, accounting, limitation, insurance and operational requirements.",
        "Duplicate working copies, downloads, email attachments and uncontrolled paper copies must be removed when no longer required for the active task.",
        "A legal, safeguarding, complaint, investigation, breach or regulator hold suspends ordinary destruction only for the records and period necessary; the reason, scope, owner and release must be documented.",
        "Processor-held data must be returned or deleted at exit and according to the approved schedule, with evidence retained.",
    ])
    add_h(doc, "2. Core retention schedule", 1)
    add_table(doc, ["Record category", "Retention rule", "Disposition"], [
        ("Unsuccessful or withdrawn admissions", "12 months after closure or waiting-list expiry unless an active dispute or documented need applies", "Secure deletion of application, uploads and working copies; retain only a justified minimal audit record."),
        ("Enrolled-pupil application", "Transfer necessary records into the pupil file; remove duplicate admissions copies after verification", "Apply the pupil academic, health, safeguarding and financial category periods."),
        ("CCTV routine footage", "30 days", "Automatic overwrite; incident exports follow the CCTV policy."),
        ("Website enquiries and unsuccessful recruitment", "12 months after closure or decision", "Delete messages, CVs, attachments and system copies unless the relationship continues."),
        ("Rights requests, complaints and breaches", "Six years proposed after closure/final report, subject to legal approval and regulator direction", "Restricted compliance archive followed by secure deletion."),
        ("Staff core employment and finance records", "Six years proposed after termination or financial year, subject to Zimbabwean employment/tax/legal review", "Restricted archive and secure destruction at expiry."),
        ("Consent records", "For the life of the relevant processing plus six years proposed for accountability", "Retain wording/version, authority, timestamp and withdrawal outcome; remove unnecessary underlying content."),
        ("Training, DPIAs, risks, audits and processor contracts", "Employment/relationship/control life plus three to six years according to accountability and claims need", "Controlled governance archive and secure destruction."),
    ], [2450, 4000, 2910], font=7.9)
    add_h(doc, "3. Children, age 18 and leavers", 1)
    add_p(doc, "At age 18 or departure, the responsible owner shall review information processed under guardian consent and delete what is no longer necessary, retain only under another recorded lawful basis and period, or obtain fresh consent directly from the adult for optional future use. Alumni, publication and fundraising records require a separate documented purpose and choice.")
    add_h(doc, "4. Secure electronic destruction", 1)
    add_p(doc, "Electronic records must be removed from live systems, shared drives, local devices, mailboxes and processors using approved deletion or secure-wipe methods. Retention in protected backups may continue only until expiry of the approved backup cycle, with ordinary access prohibited and controls preventing restored expired data from returning to active use.")
    add_h(doc, "5. Document-shredding policy", 1)
    add_p(doc, "Paper containing personal, confidential or restricted information must never be placed in ordinary waste or recycling. It must be deposited in a locked confidential-waste console or destroyed by an approved cross-cut shredder so that the information cannot practicably be reconstructed. Bulk or highly sensitive destruction shall use an approved contracted provider with sealed chain of custody, secure transport, verified destruction and a certificate of destruction. Any loss, spill or failed destruction must be treated as a possible breach.")
    add_h(doc, "6. Destruction authorisation and evidence", 1)
    add_p(doc, "Before destruction, the records owner must confirm that the retention trigger has passed and no hold applies. The School shall record the category, date range, volume, system or media, retention authority, method, date, operator, witness or contractor certificate and any exception. Destruction processes and contractor performance shall be sampled at least annually.")
    add_common_end(doc, [
        "Approved retention schedule and documented legal rationale for final category periods.",
        "System lifecycle and CCTV overwrite settings, deletion reports and processor deletion confirmations.",
        "Legal/safeguarding hold records, release decisions and age-18/leaver reviews.",
        "Confidential-waste collection logs, internal shredding records, chain-of-custody records and destruction certificates.",
    ], children_source=True)
    return save(doc, "08-Records-Retention-and-Secure-Destruction-Policy.docx")


def processor_policy():
    doc = new_policy("Third-Party Processor and Data Processing Agreement Policy", "Procurement due diligence, mandatory contracts, oversight, incidents and exit", "POL-DP-09", "Data Protection Officer, Procurement and IT")
    add_callout(doc, "Policy position", "No supplier may process personal information for Ruzawi before privacy and security due diligence is completed and a written data processing agreement imposes the required protections.")
    add_h(doc, "1. Scope", 1)
    add_p(doc, "This policy applies to school-management systems, payroll, cloud email/storage, website hosting, admissions databases, payment administration, CCTV/security providers, biometric or attendance systems, learning platforms, communications, analytics, consultants and any other supplier that handles personal information on the School's behalf.")
    add_h(doc, "2. Pre-contract assessment", 1)
    add_p(doc, "The business owner, IT and DPO shall define the purpose, data, affected people, children and sensitive categories, countries, access, retention, security, integration, subprocessors and exit requirements. A DPIA screen and children's high-risk review must be completed before trial or production data is shared. Hosting, backup, support and remote-access locations must be verified rather than assumed.")
    add_h(doc, "3. Minimum processor standards", 1)
    add_bullets(doc, [
        "Process only on documented Ruzawi instructions and only for the agreed purpose and duration.",
        "Ensure confidentiality, authorised personnel, strong identity/access, encryption where proportionate, logging, vulnerability management, continuity and tested deletion.",
        "Do not use School information for the supplier's advertising, profiling, model training or unrelated product development without a separately lawful written agreement.",
        "Support rights requests, DPIAs, audits, security testing, regulator inquiries, corrections, return and deletion.",
        "Notify Ruzawi immediately of any suspected incident and provide facts, logs and cooperation needed for the 24-hour, 72-hour and 21-day duties.",
    ])
    add_h(doc, "4. Mandatory data processing agreement", 1)
    add_p(doc, "The written contract must state subject matter, duration, nature, purpose, data types, data subjects, controller instructions, confidentiality, security, subprocessors, international transfers, rights assistance, breach cooperation, audit/information rights, retention, return/deletion, liability, insurance, termination and transition. Each authorised subprocessor must be bound by equivalent written duties, and the primary processor remains responsible for its performance.")
    add_h(doc, "5. Changes and oversight", 1)
    add_p(doc, "The contract owner shall monitor security, service, incidents, locations, subprocessors, certifications, retention and compliance throughout the relationship and at least annually for higher-risk suppliers. Material changes require DPO review before acceptance. Contract gaps, audit findings and exceptions must have owners and dates; unresolved high risk shall be escalated to the Head and Board.")
    add_h(doc, "6. Exit", 1)
    add_p(doc, "At termination, the School shall revoke access, obtain necessary data in a usable form, require deletion of live, test, support and subprocessor copies, document the backup deletion cycle and obtain written certification. The information owner must verify that integrations, accounts, keys and data flows are closed.")
    add_common_end(doc, [
        "Processor inventory and contract register identifying service, data, owner, countries, subprocessors and renewal/exit dates.",
        "Completed privacy/security due diligence, DPIA screening and children's-data assessment.",
        "Executed DPA and transfer terms, certifications, audit evidence, incident provisions and remediation.",
        "Annual supplier reviews, change approvals and end-of-contract return/deletion certification.",
    ], children_source=True)
    return save(doc, "09-Third-Party-Processor-and-Data-Processing-Agreement-Policy.docx")


def transfer_policy():
    doc = new_policy("Cross-Border Personal Information Transfer Policy", "Authority, safeguards, children's high-risk transfers and documented approval", "POL-DP-10", "Data Protection Officer")
    add_callout(doc, "Policy position", "Personal information may be transferred or remotely accessed outside Zimbabwe only after the destination, recipient, purpose, legal condition, safeguards and POTRAZ requirements have been documented and approved.")
    add_h(doc, "1. Transfers covered", 1)
    add_p(doc, "A transfer includes foreign hosting, backups, email routing, cloud storage, analytics, remote support, access by a foreign affiliate or subprocessor and disclosure to a person or organisation outside Zimbabwe. The absence of a deliberate export does not remove the need to assess provider data locations and remote access.")
    add_h(doc, "2. Approval conditions", 1)
    add_p(doc, "Before transfer, the information owner and DPO shall document the purpose and necessity, data and people affected, exporter/importer, all countries, recipient laws and practices, onward transfer, retention, security and rights. The transfer must rely on adequate protection, unambiguous consent, contractual necessity, Standard Contractual Clauses or equivalent safeguard, or another condition permitted under sections 28 and 29 of the Act. The selected condition and supporting documents must be recorded.")
    add_h(doc, "3. Children and high risk", 1)
    add_p(doc, "Transfers of children's information are treated as high risk. The School shall complete a DPIA and obtain prior POTRAZ authorisation or a documented exemption where required by CDPG 2 of 2024 before transfer begins. A parental consent clause does not replace an Authority authorisation required for high-risk children's processing.")
    add_h(doc, "4. Transfer safeguards", 1)
    add_bullets(doc, [
        "Data minimisation and, where feasible, pseudonymisation before transfer.",
        "Strong encryption in transit and at rest, with appropriate control of keys and accounts.",
        "Written processor and onward-transfer restrictions, security obligations, rights assistance and deletion duties.",
        "Assessment of recipient-country law, enforceability, government access, remedies, incident support and subprocessor locations.",
        "Clear notice to data subjects describing the transfer and available information about applicable safeguards.",
    ])
    add_h(doc, "5. Register, review and suspension", 1)
    add_p(doc, "Every transfer shall be recorded with its ROPA and DPIA references, countries, recipient, purpose, data, safeguard, POTRAZ reference, owner and review date. Transfers must be reviewed at least annually and on provider, country, subprocessor, legal, security or purpose change. The DPO may suspend a transfer where safeguards, authorisation or evidence are missing or no longer reliable.")
    add_h(doc, "6. Emergency transfers", 1)
    add_p(doc, "Any proposed emergency transfer must be necessary to protect a vital interest or meet another lawful condition, limited to the minimum information and immediately documented. The DPO and Head must be informed, and POTRAZ notification or authorisation requirements remain applicable.")
    add_common_end(doc, [
        "Current international-transfer register linked to each ROPA entry, processor and system.",
        "Recipient-country/transfer impact assessments, SCCs or equivalent safeguards and security evidence.",
        "POTRAZ requests, authorisations, exemptions, conditions and related correspondence.",
        "Data-subject notices/consents where applicable, annual reviews, change decisions and suspended-transfer records.",
    ], children_source=True)
    return save(doc, "10-Cross-Border-Personal-Information-Transfer-Policy.docx")


def dpia_policy():
    doc = new_policy("Data Protection Impact Assessment and Privacy Risk Policy", "Mandatory screening, risk assessment, authorisation, mitigation and review", "POL-DP-11", "Data Protection Officer")
    add_callout(doc, "Policy position", "No high-risk processing may begin until its necessity, proportionality, data flow, harms, controls, residual risk and POTRAZ requirements have been assessed and approved.")
    add_h(doc, "1. When a DPIA is required", 1)
    add_p(doc, "DPIA screening is mandatory before new or materially changed processing involving children, sensitive or health information, biometrics, CCTV or systematic monitoring, cross-border transfers, large-scale data, innovative technology, data matching, profiling or automated decisions, extensive sharing or another activity likely to create serious harm. Screening must occur before procurement, configuration, trial or go-live.")
    add_h(doc, "2. Assessment content", 1)
    add_bullets(doc, [
        "Specific purpose, expected benefit, necessity and less intrusive alternatives.",
        "Complete data flow from collection through use, storage, sharing, transfer, archive and deletion.",
        "Affected people, children's best interests, sensitive categories, lawful basis, notices, consent and rights.",
        "Systems, access, processors, subprocessors, countries, retention, security and incident response.",
        "Threats and potential physical, emotional, safeguarding, identity, financial, reputational, discrimination or rights harm.",
        "Inherent risk, planned controls, residual risk, owners, completion dates and monitoring measures.",
    ])
    add_h(doc, "3. Consultation", 1)
    add_p(doc, "The assessment shall consult appropriate information owners, IT/security, safeguarding, legal and suppliers. Where proportionate, the School should consider the views and reasonable expectations of pupils, parents, staff or other affected groups without exposing confidential project or security information.")
    add_h(doc, "4. Approval and POTRAZ", 1)
    add_p(doc, "The DPO shall provide written advice and determine notification or authorisation requirements. Required POTRAZ authorisation, including for applicable high-risk children's processing, must be obtained before use. The accountable owner may accept only permissible residual risk. High or critical unresolved risk must be escalated to the Head and Board and must not be accepted merely for convenience, cost or schedule.")
    add_h(doc, "5. Risk management", 1)
    add_p(doc, "Privacy risks shall be recorded with a cause, event, harm, likelihood, impact, inherent score, controls, residual score, owner and due date. Controls must be evidenced and tested. Overdue or ineffective treatment shall be escalated. Incidents, complaints, audit findings and near misses must update the risk assessment and relevant DPIA.")
    add_h(doc, "6. Review triggers", 1)
    add_p(doc, "A DPIA shall be reviewed at least annually while the high-risk processing continues and immediately upon material change to purpose, data, technology, automation, access, supplier, subprocessor, country, retention, security, law, incident or evidence of harm. Superseded versions must be retained to demonstrate decision history.")
    add_common_end(doc, [
        "DPIA screening decisions for new systems, suppliers and material changes.",
        "Completed DPIAs for CCTV, children's data, biometrics, transfers and other high-risk activities.",
        "Privacy risk register, treatment evidence, residual-risk approvals and overdue remediation escalation.",
        "POTRAZ notification/authorisation records, annual DPIA reviews and change-triggered reassessments.",
    ], children_source=True)
    return save(doc, "11-Data-Protection-Impact-Assessment-and-Privacy-Risk-Policy.docx")


def training_policy():
    doc = new_policy("Data Protection Training and Awareness Policy", "Induction, annual refreshers, role-based competence and completion evidence", "POL-DP-12", "Data Protection Officer and Human Resources")
    add_callout(doc, "Policy position", "No person should handle personal information without understanding the School's privacy, security, children, rights and incident-reporting duties. Training completion will be recorded and enforced.")
    add_h(doc, "1. Scope and frequency", 1)
    add_p(doc, "All employees, governors, regular volunteers and contractors with access to personal information must complete data protection induction before or promptly after access is granted and an annual refresher. The DPO, IT, Admissions, HR, Finance, medical/pastoral, safeguarding, CCTV/security, procurement and senior leaders require additional role-specific instruction.")
    add_h(doc, "2. Core learning", 1)
    add_bullets(doc, [
        "Personal, sensitive and children's information; lawful processing, transparency, purpose and minimisation.",
        "Guardian authority, age-appropriate information, best interests, consent and withdrawal.",
        "Recognition and immediate forwarding of rights requests in any reasonable form.",
        "MFA, phishing, secure sharing, approved systems, paper security, CCTV and document destruction.",
        "Immediate breach reporting, safe containment, evidence preservation and the 24-hour/72-hour regulatory timetable.",
        "DPIA screening, supplier/transfer approval and prohibition on unapproved systems or processing changes.",
    ])
    add_h(doc, "3. Role-based competence", 1)
    add_p(doc, "Personnel responsible for higher-risk activities must demonstrate practical competence through scenarios or supervised tasks. Admissions must understand guardian verification and sensitive consent; DPO and information owners must handle rights requests; incident leaders must exercise breach response; Security must operate CCTV and retention; Procurement/IT must apply processor, transfer and DPIA gates; the Board must understand oversight and accountability.")
    add_h(doc, "4. Delivery and assessment", 1)
    add_p(doc, "Training may be delivered in person, online or through approved external providers, but must use the current policy and legal framework. Understanding shall be assessed through knowledge checks, scenarios, exercises or observed practice appropriate to the role. A failed assessment or material policy breach requires remedial training and may result in temporary access restriction.")
    add_h(doc, "5. Completion and escalation", 1)
    add_p(doc, "HR and the DPO shall maintain a central record of audience, module/version, date, method, result, evidence and refresher due date. Completion is mandatory. Overdue personnel shall receive reminders and escalation to their manager; continued non-compliance shall be escalated to the Head and may result in access suspension.")
    add_h(doc, "6. Awareness programme", 1)
    add_p(doc, "Formal training shall be supported by proportionate ongoing awareness, including privacy reminders, phishing or security exercises, quick breach-reporting guidance, policy updates and targeted communications after incidents or control changes. The DPO shall evaluate recurring errors, incidents, requests and audit findings to adjust the programme.")
    add_h(doc, "7. Annual exercise and Board briefing", 1)
    add_p(doc, "At least annually, Ruzawi shall conduct a breach tabletop exercise involving DPO, IT, leadership and relevant information owners and provide a compliance briefing to the Board. Actions, owners and completion evidence must be retained.")
    add_common_end(doc, [
        "Current curriculum and training materials linked to approved policy versions.",
        "Central completion records, assessment results, certificates and overdue escalation.",
        "Induction evidence and role-specific competence records for higher-risk functions.",
        "Awareness campaigns, phishing results, annual breach exercise and Board briefing minutes.",
    ], children_source=True)
    return save(doc, "12-Data-Protection-Training-and-Awareness-Policy.docx")


def main():
    builders = [
        governance_policy, lawful_processing_policy, rights_policy, children_policy,
        security_policy, cctv_policy, breach_policy, retention_policy,
        processor_policy, transfer_policy, dpia_policy, training_policy,
    ]
    paths = [builder() for builder in builders]
    print("\n".join(str(path) for path in paths))


if __name__ == "__main__":
    main()
