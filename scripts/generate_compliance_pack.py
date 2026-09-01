from pathlib import Path
from datetime import date
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "data-protection-compliance-pack"
LOGO = ROOT / "public" / "images" / "ruzawi-logo.png"
TODAY = "21 August 2026"
REVIEW = "21 August 2027"

GREEN = "00582C"
BLUE = "47778D"
PALE_BLUE = "E8F1F5"
PALE_GREEN = "EAF3ED"
GOLD = "B89A50"
INK = "10251C"
MID = "4E5C53"
LIGHT = "F3F5F3"
WHITE = "FFFFFF"
RED = "8B1E2D"
AMBER = "8A6500"


def set_run_font(run, size=None, bold=None, italic=None, color=INK, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        h = tr_pr.find(qn("w:trHeight"))
        if h is not None:
            tr_pr.remove(h)
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("Page ")
    set_run_font(r, 8.5, color=MID)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    rr = paragraph.add_run()._r
    rr.extend([fld_begin, instr, fld_sep, txt, fld_end])


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GREEN)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.16

    for name, size, color, before, after in (
        ("Title", 27, GREEN, 0, 5),
        ("Subtitle", 12.5, MID, 0, 18),
        ("Heading 1", 16, GREEN, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, GREEN, 9, 4),
    ):
        s = styles[name]
        s.font.name = "Aptos Display" if name != "Subtitle" else "Aptos"
        s._element.rPr.rFonts.set(qn("w:ascii"), s.font.name)
        s._element.rPr.rFonts.set(qn("w:hAnsi"), s.font.name)
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = name != "Subtitle"
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
        s.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        s = styles[name]
        s.font.name = "Aptos"
        s.font.size = Pt(10.5)
        s.paragraph_format.left_indent = Inches(0.38)
        s.paragraph_format.first_line_indent = Inches(-0.19)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.16

    if "Small Text" not in [s.name for s in styles]:
        s = styles.add_style("Small Text", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        s.font.name = "Aptos"
        s.font.size = Pt(8.5)
        s.font.color.rgb = RGBColor.from_string(MID)
        s.paragraph_format.space_after = Pt(3)

    if "Form Prompt" not in [s.name for s in styles]:
        s = styles.add_style("Form Prompt", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        s.font.name = "Aptos"
        s.font.size = Pt(10)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(GREEN)
        s.paragraph_format.space_before = Pt(5)
        s.paragraph_format.space_after = Pt(2)
        s.paragraph_format.keep_with_next = True


def new_doc(title, subtitle, doc_id, owner="Data Protection Officer"):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.32)
    sec.footer_distance = Inches(0.35)
    setup_styles(doc)

    header = sec.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"RUZAWI SCHOOL  |  {doc_id}")
    set_run_font(r, 8, bold=True, color=BLUE)

    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Controlled document - draft for approval and operationalisation")
    set_run_font(r, 8, color=MID)
    add_page_field(footer.add_paragraph())

    if LOGO.exists():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        image = p.add_run().add_picture(str(LOGO), width=Inches(0.75))
        image._inline.docPr.set("descr", "Ruzawi School crest")
        image._inline.docPr.set("title", "Ruzawi School")
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    r = kicker.add_run("DATA PROTECTION COMPLIANCE PACK")
    set_run_font(r, 9, bold=True, color=GOLD)
    t = doc.add_paragraph(title, style="Title")
    t.paragraph_format.keep_with_next = True
    s = doc.add_paragraph(subtitle, style="Subtitle")
    s.paragraph_format.keep_with_next = True

    table = doc.add_table(rows=5, cols=2)
    set_table_geometry(table, [2100, 7260])
    rows = [
        ("Document ID", doc_id),
        ("Status", "DRAFT - requires approval, named owners and implementation evidence"),
        ("Owner", owner),
        ("Effective / review", f"Effective on approval / review no later than {REVIEW}"),
        ("Approval", "Head and/or Board: [NAME / SIGNATURE / DATE]"),
    ]
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        shade(table.cell(i, 0), PALE_GREEN)
        for run in table.cell(i, 0).paragraphs[0].runs:
            set_run_font(run, 9, bold=True, color=GREEN)
        for run in table.cell(i, 1).paragraphs[0].runs:
            set_run_font(run, 9, color=INK)
    doc.add_paragraph()
    return doc


def add_h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_p(doc, text="", bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, 10.5, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, 10.5, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, 10.5, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, 10.5)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, 10.5)


def add_callout(doc, label, text, kind="info"):
    color = GREEN if kind == "info" else (AMBER if kind == "warning" else RED)
    fill = PALE_GREEN if kind == "info" else ("FFF4D6" if kind == "warning" else "FBEAEC")
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label.upper() + "  ")
    set_run_font(r, 9, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, 9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, font=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    repeat_header(table.rows[0])
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        shade(table.cell(0, i), GREEN)
        for r in table.cell(0, i).paragraphs[0].runs:
            set_run_font(r, font, bold=True, color=WHITE)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            if row_idx % 2:
                shade(cells[i], LIGHT)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.05
                for r in p.runs:
                    set_run_font(r, font, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def form_line(doc, label, hint=""):
    p = doc.add_paragraph(style="Form Prompt")
    p.add_run(label)
    p2 = doc.add_paragraph("________________________________________________________________________________")
    p2.paragraph_format.space_after = Pt(4)
    for r in p2.runs:
        set_run_font(r, 9, color=BLUE)
    if hint:
        p3 = doc.add_paragraph(hint, style="Small Text")
        p3.paragraph_format.space_after = Pt(4)


def add_sources(doc, extra=None, compact_review=False):
    add_h(doc, "Legal and reference sources", 1)
    sources = [
        ("Cyber and Data Protection Act [Chapter 12:07]", "https://www.potraz.gov.zw/wp-content/uploads/2026/02/ACT-CDPA.pdf"),
        ("Cyber and Data Protection (Licensing of Data Controllers and Appointment of Data Protection Officers) Regulations, 2024 (S.I. 155 of 2024)", "https://www.potraz.gov.zw/wp-content/uploads/2025/02/sI-155-of-2024-Cyber-and-Data-Protection-Normal_240913_1250178.pdf"),
        ("POTRAZ Data Protection Authority", "https://www.potraz.gov.zw/"),
    ]
    if extra:
        sources.extend(extra)
    for name, url in sources:
        p = doc.add_paragraph(style="Small Text")
        add_hyperlink(p, name, url)
    p = doc.add_paragraph(style="Small Text")
    p.add_run("Checklist source: ATS-CYBERNESIS Schools Data Protection Compliance Checklist V3, derived from CDPG 1 of 2025 and incorporating CDPG 2 of 2024.")
    review_text = "These documents are operational templates, not legal advice or proof of compliance. Ruzawi must validate them against its actual systems, contracts, licences and practices, obtain appropriate Zimbabwean legal/DPO review, approve them, train staff, and retain implementation evidence."
    if compact_review:
        p = doc.add_paragraph(style="Small Text")
        r = p.add_run("REVIEW NOTE  ")
        set_run_font(r, 8.5, bold=True, color=AMBER)
        r = p.add_run(review_text)
        set_run_font(r, 8.5, color=INK)
    else:
        add_callout(doc, "Review note", review_text, "warning")


def save(doc, name):
    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / name
    doc.core_properties.title = name.replace(".docx", "")
    doc.core_properties.subject = "Ruzawi School data protection compliance pack"
    doc.core_properties.author = "Ruzawi School"
    doc.core_properties.keywords = "data protection, POTRAZ, CDPA, compliance"
    doc.save(path)
    return path


def build_status_report():
    doc = new_doc(
        "Implementation Status and Evidence Matrix",
        "Repository-based readiness review against the ATS-CYBERNESIS school checklist",
        "DP-00",
    )
    add_callout(doc, "Conclusion", "The public-facing website privacy layer is materially developed, but full compliance cannot be confirmed. Critical regulatory, contractual, technical, physical and operational evidence remains unverified. The documents in this pack close documentation gaps only after approval and implementation.", "warning")
    add_h(doc, "1. Scope and assurance level", 1)
    add_p(doc, "Reviewed evidence: the supplied checklist; the website Privacy Policy, Cookie Policy, Terms of Use, privacy-request page and handling procedure; admissions declarations and server-side validation; CookieHub consent gating; and repository code relevant to forms and privacy.")
    add_p(doc, "Not reviewed: POTRAZ licence/DP1/DP2/DP3 records, Board or SMT minutes, HR/personnel files, contracts and processor terms, cloud account settings, security logs, CCTV configuration, physical records controls, training records, actual requests, actual incidents, backups, disposal records, DPIAs or regulator correspondence.")
    add_h(doc, "2. What is demonstrably in place", 1)
    add_bullets(doc, [
        "A detailed public Privacy Policy naming Ruzawi as controller, describing purposes, recipients, cross-border processing, retention, security, rights, children, consent and contact routes.",
        "A public Cookie Policy and CookieHub integration; Google Analytics and Microsoft Clarity tags are marked for consent-based loading.",
        "A dedicated privacy-request form, monitored-address requirement and a documented internal request-handling procedure with identity/authority safeguards and a 30-day control target.",
        "Admissions-specific transparency and separate electronic declarations for guardian authority, sensitive data, international transfer, third-party authority, privacy-notice acknowledgement and terms.",
        "Server-side validation of required declarations, notice version and timestamps before an application is accepted, plus reCAPTCHA protection.",
        "A public statement that admissions decisions are not made solely by automated processing.",
    ])
    add_h(doc, "3. Priority gaps before claiming compliance", 1)
    add_table(doc, ["Priority", "Gap", "Required evidence / action"], [
        ("Critical", "Controller registration and DPO status not evidenced", "Obtain/locate current DP1, licence certificate, fee receipt, DPO appointment, qualification/certification and DP2 acknowledgement."),
        ("Critical", "Children's high-risk processing and transfers", "Inventory CCTV, cloud hosting, backups, analytics, reCAPTCHA, email, SIS and other children's-data transfers; complete DPIAs; obtain POTRAZ authorisation/exemption where required by the children's guideline."),
        ("Critical", "Processor contracts and transfer safeguards", "Execute DPAs; record hosting locations, subprocessors and transfer mechanism; complete transfer impact reviews; retain SCCs/IDTAs or equivalent and Authority correspondence."),
        ("High", "Operational security evidence", "Verify MFA, least privilege, encryption, logging, backup restoration tests, patching, malware controls, CCTV access/retention and physical records security."),
        ("High", "Governance and approval", "Name owners, approve and sign policies, record Board/SMT commitment, budget resources, communicate policies and maintain version control."),
        ("High", "Training and incident readiness", "Train all staff and alternates; retain completion records; conduct a breach tabletop exercise and close actions."),
        ("High", "Registers and implementation records", "Populate ROPA, consent, DSAR, breach, transfer, DPIA, risk, disposal, training and vendor registers with real evidence."),
        ("High", "Guardianship verification", "Do not rely only on applicant self-declaration. Implement a proportionate evidence/sighting step for parental or legal-guardian authority and record the outcome."),
        ("Medium", "Age-appropriate child information", "Issue the child-facing privacy explanation in DP-04 and record DPO review and communication."),
        ("Medium", "Retention precision", "Approve category-specific periods after legal and operational review, configure deletion/overwrite, and retain destruction evidence."),
    ], [1000, 2900, 5460], font=8.2)
    add_h(doc, "4. Checklist evidence matrix", 1)
    rows = [
        (1,"DPO appointment","Not verified","DP-02 supplies appointment and role templates; obtain DP2 acknowledgement and qualification evidence."),
        (2,"Registration/licence","Not verified","Locate current DP1, licence certificate, tier decision, receipt and renewal reminder."),
        (3,"Roles and responsibilities","Draft supplied","Approve DP-02; insert names and organogram; update job descriptions."),
        (4,"Approved Data Protection Policy","Partial","Public notice exists; approve and sign DP-01 and issue the child notice."),
        (5,"Senior management commitment","Not verified","Record Board/SMT approval, budget and review minutes."),
        (6,"Lawful basis by processing category","Partial","Public notice describes bases; validate and populate DP-03 ROPA."),
        (7,"Specific/informed consent","Strong website evidence","Validate wording with DPO/counsel; retain consent records and guardianship evidence."),
        (8,"Consent management","Partial","Application timestamps exist; implement central register and withdrawal workflow in DP-04/DP-12."),
        (9,"Purpose limitation","Partial","Policy statement exists; evidence access roles, periodic reviews and change control."),
        (10,"Processor agreements","Not verified","Execute DPAs and complete DP-09 records for each provider."),
        (11,"Rights procedure","Documented","Existing repository SOP strengthened in DP-05; approve and train staff."),
        (12,"Requests within 30 days","Not verified","Populate DSAR register and review overdue metrics."),
        (13,"Rights request records","Not verified","Use DP-12 register; retain redacted closed examples."),
        (14,"Security measures","Partial","Code shows consent gating and reCAPTCHA; operational IT/physical controls require evidence under DP-07."),
        (15,"Incident/breach procedure","Draft supplied","Approve DP-06, publish internal contact route and exercise it."),
        (16,"Breach training","Not verified","Deliver DP-11 modules and retain completion records."),
        (17,"Regulator/data-subject breach reporting","Not verified","Use DP-06/DP-12; retain DP3, 24/72-hour evidence and 21-day report."),
        (18,"Cross-border authorisation","Not verified","High-priority review of website/cloud providers and POTRAZ authorisation."),
        (19,"Transfer safeguards","Not verified","Retain SCCs/equivalent, DPA, security certification and impact review."),
        (20,"People informed of transfers","Documented online","Privacy and admissions notices disclose overseas processing; validate accuracy and provider countries."),
        (21,"Data minimisation","Partial","Forms label optional fields and warn against excess data; perform formal field-by-field audit."),
        (22,"Retention periods","Draft supplied","Public notice is principle-based; approve DP-08 periods and configure systems."),
        (23,"Secure disposal","Draft supplied","Implement destruction/wiping process and logs."),
        (24,"Initial/refresher training","Not verified","Deliver annual and role-based training."),
        (25,"Training completion records","Not verified","Use DP-11/DP-12 register and certificates."),
        (26,"Induction awareness","Draft supplied","Insert module into official induction checklist and retain sign-off."),
        (27,"DPIAs for high risk","Not verified","Screen all systems and complete DPIAs for CCTV, applications, cloud/SIS/biometrics as applicable."),
        (28,"DPIA review","Not verified","Populate review dates, triggers and DPO approvals."),
        (29,"Privacy risk management","Draft supplied","Validate initial risk register, assign owners and evidence mitigation."),
        (30,"Parental consent and authority verification","Partial","Electronic declaration exists; implement documentary verification/sighting record."),
        (31,"Age-appropriate notices","Gap / draft supplied","Issue and explain the child-facing notice in DP-04."),
        (32,"Best interests / privacy by design","Partial","Website minimisation and notices help; apply DPIA/vendor gate and pupil e-safety guidance school-wide."),
        (33,"Prior authorisation for high-risk child processing","Not verified","Classify and seek POTRAZ authorisation or documented exemption before use."),
        (34,"POTRAZ notice for non-consent child processing","Not verified","Document bases and notifications; obtain legal advice on each applicable activity."),
        (35,"No solely automated significant decisions","Partial","Admissions statement exists; verify all pupil systems and document human review."),
        (36,"Age-18/leaver review","Draft supplied","Configure trigger and record deletion, retention or refreshed consent."),
    ]
    add_table(doc, ["Item","Requirement","Repository assessment","Next evidence"], rows, [600, 2200, 1800, 4760], font=7.5)
    add_h(doc, "5. Approval and implementation sequence", 1)
    add_numbers(doc, [
        "Confirm the legal entity name, licence tier, DPO and privacy governance owners.",
        "Complete the ROPA and provider/transfer inventory using actual contracts and system settings.",
        "Prioritise DPIAs and POTRAZ engagement for children's high-risk processing and cross-border transfers.",
        "Approve the policy and procedure documents; insert names, contacts, signatures and dates.",
        "Configure technical, physical, retention and access controls, then capture screenshots/logs as evidence.",
        "Train all staff; run a breach tabletop exercise and remediate findings.",
        "Populate registers continuously and perform a quarterly DPO review plus annual Board report.",
    ])
    add_sources(doc)
    return save(doc, "00-Implementation-Status-and-Evidence-Matrix.docx")


def build_policy():
    doc = new_doc("Data Protection Policy", "School-wide rules for lawful, fair, secure and accountable processing", "DP-01")
    add_callout(doc, "Policy status", "This internal policy complements, but does not replace, the public website Privacy Policy. It becomes evidence only when approved, communicated and implemented.")
    add_h(doc, "1. Purpose and scope", 1)
    add_p(doc, "Ruzawi School is committed to respecting the privacy and dignity of pupils, parents and guardians, staff, applicants, alumni, visitors, suppliers and other data subjects. This policy applies to all personal information processed in paper or electronic form by employees, governors, contractors, volunteers and processors acting for the School.")
    add_h(doc, "2. Governing principles", 1)
    add_bullets(doc, [
        "Lawfulness, fairness and transparency: identify and record a valid basis; give clear information at collection or as soon as required.",
        "Purpose limitation: use information only for specified, explicit and legitimate purposes; assess and authorise any new incompatible use.",
        "Data minimisation: collect and expose only what is adequate, relevant and necessary.",
        "Accuracy: enable correction and take reasonable steps to keep material information current.",
        "Storage limitation: apply approved retention periods and age/leaver review triggers; securely delete, destroy or anonymise at expiry.",
        "Security and confidentiality: use risk-based technical, organisational and physical controls through the full data life cycle.",
        "Accountability: maintain records capable of demonstrating compliance to data subjects and POTRAZ.",
        "Best interests of the child: design processing to protect children by default and avoid unnecessary exposure, profiling or over-sharing.",
    ])
    add_h(doc, "3. Lawful basis and special categories", 1)
    add_p(doc, "The responsible owner must record the purpose, categories, data subjects, recipients, retention, location and lawful basis in the ROPA before processing begins. Consent must be specific, informed, unambiguous, demonstrable, separable from unrelated terms and as easy to withdraw as to give. Written consent is required for sensitive information unless a documented statutory exception applies.")
    add_p(doc, "For a child, the School must obtain written consent from a parent or legal guardian where consent is the basis, verify that person's authority using proportionate evidence, and provide information in language that the child and guardian can understand. Any non-consent basis, high-risk processing, surveillance/tracking, cross-border transfer or significant automated decision involving a child must pass DPO/legal review and applicable POTRAZ notification or authorisation steps before use.")
    add_h(doc, "4. Transparency and data subject rights", 1)
    add_bullets(doc, [
        "Issue an appropriate privacy notice covering controller identity and address, purposes, required/optional fields and consequences, recipients, transfers, rights and contact route.",
        "Forward all rights requests immediately to privacy@ruzawi.com; do not discourage or narrow a request before DPO review.",
        "Authenticate proportionately, verify representative authority, search all relevant systems, protect third-party information and document the decision.",
        "Use 30 calendar days as the internal completion control unless the DPO records a different legally applicable period or regulator direction.",
        "Permit withdrawal of consent at any time, without explanation and free of charge; stop future consent-based processing unless another valid basis is documented.",
    ])
    add_h(doc, "5. Children and automated decisions", 1)
    add_p(doc, "Children must not be subject to a decision based solely on automated processing that produces legal or similarly significant effects unless expressly permitted by law and approved by the DPO with meaningful human intervention. Pupil-facing systems must use privacy-protective defaults, age-appropriate explanations and controls discouraging over-sharing.")
    add_h(doc, "6. Processors and international transfers", 1)
    add_p(doc, "No processor may receive personal information before documented due diligence, a written data processing agreement and approval by the information owner and DPO. The School must maintain a processor/contract register and international-transfer register. Cross-border transfers require the safeguards, consent, authorisation or other legal condition applicable under the Act and POTRAZ guidance; transfers involving children's information must receive enhanced review and any required prior Authority authorisation.")
    add_h(doc, "7. Security, incidents and continuity", 1)
    add_bullets(doc, [
        "Apply least privilege, unique accounts, MFA where available, encryption in transit and at rest where proportionate, secure configuration, patching, endpoint protection, logging and monitored backups.",
        "Protect paper records, server/network areas and CCTV with restricted access, visitor controls and documented retention.",
        "Report suspected loss, misdirection, unauthorised access, disclosure, alteration or unavailability immediately using the breach plan; preserve evidence and do not conceal incidents.",
        "Notify POTRAZ within 24 hours of awareness of a personal data breach; use Form DP3; notify affected people within 72 hours where the breach is likely to create high risk; conclude and submit the final investigation report within 21 days of notification.",
    ])
    add_h(doc, "8. DPIAs and change control", 1)
    add_p(doc, "A DPIA screening is mandatory before new or materially changed processing involving children, sensitive or biometric information, CCTV/surveillance, large-scale data, profiling/automation, innovative technology, systematic monitoring or international transfer. High-risk processing may not go live until risks are reduced, the DPO approves the outcome and required POTRAZ engagement is complete.")
    add_h(doc, "9. Training, monitoring and enforcement", 1)
    add_p(doc, "All staff and regular contractors must complete induction and annual refresher training. Higher-risk roles receive role-specific training. The DPO shall conduct quarterly register reviews, an annual self-assessment/internal audit and an annual report to the Head/Board. Breach of this policy may result in access suspension, disciplinary action, contract remedies, regulator notification or legal action.")
    add_h(doc, "10. Responsibilities", 1)
    add_table(doc, ["Role","Core accountability"], [
        ("Board / Head","Approve policy; resource compliance; review material risks and incidents; evidence oversight."),
        ("DPO","Monitor compliance, advise, train, audit, handle requests, oversee DPIAs, liaise with POTRAZ and data subjects."),
        ("Information owners","Maintain accurate ROPA entries, access approvals, retention and risk controls for their functions."),
        ("IT / security","Implement, test and evidence technical controls; investigate and contain incidents."),
        ("HR / Admissions / Finance / Pastoral / Medical","Apply notices, lawful basis, consent, minimisation, access, retention and rights procedures."),
        ("All personnel","Use information only as authorised; secure it; report incidents and rights requests immediately."),
    ], [2300, 7060], font=8.7)
    add_sources(doc)
    return save(doc, "01-Data-Protection-Policy.docx")


def build_governance():
    doc = new_doc("Governance, Roles and DPO Appointment Pack", "Appointment instruments, reporting lines and evidence controls", "DP-02", "Head / Board")
    add_h(doc, "1. Governance model", 1)
    add_table(doc, ["Role","Named person","Reports to","Evidence required"], [
        ("Accountable executive / Head","[NAME]","Board","Approval minutes; budget; annual review."),
        ("Data Protection Officer","[NAME / QUALIFICATION]","Head and Board","Signed appointment; certification/training; DP2 acknowledgement."),
        ("Deputy / Privacy Champion","[NAME]","DPO","Delegation; training; mailbox cover."),
        ("IT Security Lead","[NAME]","Head / DPO","Control reports; incident response records."),
        ("Information owners","Admissions / HR / Finance / Pastoral / Medical / Boarding / IT","Head / DPO","ROPA, access reviews, retention and risk entries."),
    ], [2100, 2100, 1800, 3360], font=8.2)
    add_h(doc, "2. DPO appointment letter", 1)
    add_p(doc, "To: [FULL NAME]")
    add_p(doc, "Date: [DATE]")
    add_p(doc, "Appointment as Data Protection Officer")
    add_p(doc, "Ruzawi School appoints you as its Data Protection Officer with effect from [DATE], subject to the Cyber and Data Protection Act [Chapter 12:07], S.I. 155 of 2024, applicable POTRAZ guidance and this appointment instrument.")
    add_p(doc, "You shall have timely access to the Head and Board, adequate independence and resources, and authority to obtain information needed to discharge your role. You must promptly declare conflicts and maintain confidentiality.")
    add_bullets(doc, [
        "monitor the Act, regulations, policies and data protection programme; manage internal activities and audits;",
        "raise awareness and deliver or coordinate staff training;",
        "deal with requests from POTRAZ and data subjects; advise personnel on their obligations;",
        "advise on and monitor DPIAs, high-risk children's processing and transfers;",
        "act as the contact point for POTRAZ and data subjects; maintain required registers and report material risk to the Head/Board;",
        "ensure material breaches are assessed and notifications/investigations completed within required timelines.",
    ])
    add_p(doc, "Accepted by DPO: __________________  Date: __________")
    add_p(doc, "Approved for Ruzawi School: __________________  Role: __________  Date: __________")
    add_h(doc, "3. Board / SMT approval minute template", 1)
    add_p(doc, "Resolved that the Board/SMT: (a) approves the Ruzawi Data Protection Compliance Pack version [VERSION]; (b) appoints [NAME] as DPO; (c) authorises submission/maintenance of Forms DP1 and DP2 and associated fees; (d) directs completion of the ROPA, DPIAs, vendor contracts and training plan; and (e) requires quarterly DPO reporting and immediate escalation of material breaches or high residual risks.")
    form_line(doc, "Meeting, date and minute reference")
    form_line(doc, "Attendees and conflicts declared")
    form_line(doc, "Budget/resources approved")
    form_line(doc, "Chair signature and date")
    add_h(doc, "4. Incident escalation matrix", 1)
    add_table(doc, ["Event","Immediate recipient","Escalation"], [
        ("Suspected rights request","DPO / privacy mailbox","DPO logs and assigns; owner supports search."),
        ("Suspected personal data incident","DPO + IT Security + relevant owner","Head immediately if material; Board Chair for severe/high-risk breach."),
        ("New high-risk system / transfer","DPO before procurement or go-live","Head/Board and POTRAZ where notification/authorisation is required."),
        ("Overdue remediation / repeated control failure","DPO","Head; Board/Audit Committee if unresolved."),
        ("Regulator contact","DPO","Head and legal adviser; preserve complete correspondence."),
    ], [3000, 2800, 3560], font=8.5)
    add_h(doc, "5. Quarterly DPO report minimum content", 1)
    add_bullets(doc, [
        "Licence/registration status and upcoming deadlines; DPO certification and resourcing.",
        "Rights requests received, closed, overdue and themes; consent withdrawals.",
        "Incidents, notifications, investigation actions and lessons learned.",
        "DPIAs, transfers, processor due diligence and contract gaps.",
        "Training completion by function; overdue staff; phishing/awareness results.",
        "Risk-register changes, high residual risks, audit findings and remediation ageing.",
    ])
    add_sources(doc)
    return save(doc, "02-Governance-Roles-and-DPO-Appointment-Pack.docx")


def build_ropa():
    doc = new_doc("Record of Processing Activities and Lawful Basis Register", "Initial website-based inventory for validation and expansion", "DP-03")
    add_callout(doc, "Important", "These are initial entries inferred from the website repository. Information owners must validate actual systems, countries, recipients, volumes, safeguards, retention and legal bases; add offline and school-wide processing before approval.", "warning")
    add_h(doc, "1. ROPA maintenance rules", 1)
    add_bullets(doc, [
        "Create one entry for each distinct purpose, not merely each application or department.",
        "Record direct and indirect collection, special-category/children's information, processors, countries, access roles, retention and security measures.",
        "Record the legal basis and any consent, legitimate-interest, vital-interest, contract, legal-obligation or POTRAZ approval evidence.",
        "Review at least quarterly and before any new purpose, provider, transfer, field, automated feature or material system change.",
    ])
    add_h(doc, "2. Initial processing inventory", 1)
    rows = [
        ("P-001","Online admissions","Prospective pupils; guardians; siblings; fee payer; referees","Identity/contact; photos; education; health/religion/family/financial; proof of payment","Consent; written consent for sensitive data; pre-contract steps/other basis to validate","Netlify; application database; Resend; reCAPTCHA; authorised school staff","Outside Zimbabwe likely - validate countries, safeguards and POTRAZ authorisation","Application outcome + approved waiting-list/accounting/legal periods; define in DP-08"),
        ("P-002","Privacy rights requests","Requesters; children via guardian; representatives","Identity/contact; request details; authority; verification outcome; correspondence","Legal obligation/accountability; consent only where genuinely optional","Netlify; Resend/email; DPO/privacy team; relevant information owners","Potential email/hosting transfer - validate","Request closure + dispute/audit period; propose 6 years after closure, legal review"),
        ("P-003","Contact enquiries","Website visitors; parents; prospective families","Name; email/phone; enquiry; security logs","Consent/initiated request; legitimate interest in responding","Netlify; Resend; relevant school office; reCAPTCHA","Potential hosting/email transfer - validate","12 months after closure unless converted to relationship or legal hold"),
        ("P-004","Junior Master/Mistress applications","Adult applicants; referees","Identity/contact; CV; qualifications; employment history; messages","Pre-contract steps; consent; legitimate interest; employment-law basis as applicable","Netlify; email/Resend; recruitment staff; reCAPTCHA","Potential hosting/email transfer - validate","Unsuccessful candidates: propose 12 months; successful: personnel schedule"),
        ("P-005","Website security and delivery","Visitors/users","IP; browser/device; logs; form telemetry","Legitimate interest / security obligations","Netlify; CookieHub; Google reCAPTCHA","Potential cross-border","Security logs: define proportionately, commonly 30-180 days"),
        ("P-006","Consented analytics","Website visitors who consent","Identifiers; page/event; device/browser; approximate location; replay/heatmap data","Consent","Google Analytics; Microsoft Clarity; CookieHub","Outside Zimbabwe / vendor locations","Vendor/tool configuration; validate exact cookie and account retention"),
    ]
    add_table(doc, ["ID","Purpose","Data subjects","Data","Basis","Recipients / systems","Transfer","Retention"], rows, [600,1100,1150,1500,1450,1400,1100,1060], font=6.7)
    add_h(doc, "3. Mandatory school-wide additions", 1)
    add_bullets(doc, [
        "Current pupil administration, teaching, assessment, reports and learning support.",
        "Safeguarding, pastoral care, boarding, health/sanatorium and emergency response.",
        "Fees, accounting, banking, debt management, audit and tax/NSSA/ZIMRA obligations.",
        "Staff, governor, volunteer and contractor recruitment, HR, payroll, performance, discipline and occupational health.",
        "CCTV, access control, visitor management, transport, trips, sport, clubs, chapel/religious observance and photographs/publications.",
        "School email, cloud document storage, backups, devices, cybersecurity monitoring and disaster recovery.",
        "Alumni/ROPA relations, fundraising, events, magazines, archives and direct marketing.",
        "Suppliers, procurement, legal claims, insurance, regulator correspondence and incident records.",
    ])
    add_h(doc, "4. Blank ROPA entry", 1)
    for label, hint in [
        ("Processing ID, owner and review date", "Assign a stable ID; name the accountable information owner and DPO reviewer."),
        ("Purpose and necessity", "State a specific purpose and why the processing is necessary."),
        ("Data subjects and personal-information categories", "Identify children and sensitive/health/biometric/genetic information separately."),
        ("Source, collection method and privacy notice", "Direct/indirect; form/system; notice version and date."),
        ("Lawful basis and evidence", "Record consent record, law, contract, vital/public/legitimate interest or other basis; include balancing test where used."),
        ("Systems, locations, recipients and processors", "Include paper records, internal access roles, vendors and subprocessors."),
        ("International transfer", "Country, adequacy/safeguard/consent, transfer impact review and POTRAZ authorisation reference."),
        ("Retention and disposal", "Period/trigger, legal rationale, deletion/overwrite method and owner."),
        ("Security and DPIA", "Controls, risk rating, DPIA ID, residual risk and approvals."),
    ]:
        form_line(doc, label, hint)
    add_sources(doc)
    return save(doc, "03-ROPA-and-Lawful-Basis-Register.docx")


def build_consent_children():
    doc = new_doc("Consent and Children's Data Procedure", "Consent lifecycle, guardian verification and best-interests controls", "DP-04")
    add_h(doc, "1. When this procedure applies", 1)
    add_p(doc, "Apply this procedure whenever Ruzawi relies on consent, processes sensitive information, or processes a child's personal information. A child is under 18. The DPO must be consulted where another legal basis is proposed, a child is affected by high-risk processing, or consent cannot reasonably be obtained.")
    add_h(doc, "2. Consent standard", 1)
    add_bullets(doc, [
        "Explain the controller, purpose, data, recipients, transfers, consequences of refusal, retention, rights and withdrawal route before consent.",
        "Use plain language, separate purposes and unticked opt-in choices; do not bundle optional consent with enrolment or employment where it is not necessary.",
        "For sensitive information, capture written consent and preserve the exact wording/version, person, authority, date, method and timestamp.",
        "Make withdrawal free and as easy as giving consent; confirm the outcome and stop future consent-based processing promptly.",
        "Refresh consent when purpose, data, recipient, risk or transfer materially changes and review consent given for a child at age 18 or earlier if legal capacity changes.",
    ])
    add_h(doc, "3. Guardian authority verification", 1)
    add_numbers(doc, [
        "Record the consenting adult's full name, relationship/capacity, contact details and the child concerned.",
        "Sight an appropriate birth certificate, adoption order, custody order, guardianship order or other reliable evidence. Do not retain a full copy unless necessary and authorised; otherwise record document type, date sighted, verifier and result.",
        "Resolve conflicting guardianship instructions or restrictions with the DPO and appropriate safeguarding/legal personnel before disclosure or processing.",
        "Retain the verification record with the consent record and review when notified of changes.",
    ])
    add_h(doc, "4. High-risk and non-consent child processing gate", 1)
    add_table(doc, ["Question","If yes"], [
        ("New surveillance, tracking, CCTV, biometrics or innovative pupil technology?","Complete DPIA and classify for prior POTRAZ authorisation before use."),
        ("Children's information transferred or accessed outside Zimbabwe?","Complete transfer inventory/DPIA and seek required prior POTRAZ authorisation or documented exemption; do not rely on form consent alone."),
        ("Could processing create discrimination, reputational, financial, identity-theft or significant privacy harm?","DPIA, risk reduction, DPO/Head sign-off and Authority engagement before use."),
        ("Processing a child's data without consent?","Document the exact lawful basis and best-interests analysis; notify POTRAZ in writing where required by CDPG 2 of 2024."),
        ("Automated recommendation or decision affects placement, discipline, access or another significant outcome?","Ensure meaningful human review, ability to challenge and clear explanation; prohibit solely automated decision unless legally permitted."),
    ], [4750, 4610], font=8.3)
    add_h(doc, "5. Consent withdrawal workflow", 1)
    add_numbers(doc, [
        "Log request, verify identity/authority proportionately and acknowledge receipt.",
        "Identify every system, processor and recipient using the relevant consent.",
        "Stop future processing under that consent; determine whether a separate lawful basis requires limited continuation and document it.",
        "Instruct processors, amend preferences, suppress future use and record technical completion.",
        "Confirm outcome, explain any retained information and close the register entry after DPO review.",
    ])
    add_h(doc, "6. Guardian verification record", 1)
    for label in ["Child name / pupil ID", "Adult name and capacity", "Evidence type and reference (avoid unnecessary copying)", "Date sighted; sighted by", "Restrictions or conflicts identified", "Verification result and DPO escalation", "Signature / timestamp"]:
        form_line(doc, label)
    add_h(doc, "7. Child-friendly privacy notice", 1)
    add_callout(doc, "For pupils", "Your information belongs to you. Ruzawi uses it to teach you, keep you safe and healthy, organise school life and communicate with your family. We only ask for information we need. Some trusted companies help us run school systems, and we must make them protect it. Ask a trusted adult or the privacy team if you want to know what we hold, correct something wrong, or are worried about how it is used.")
    add_p(doc, "We will not make an important decision about you only by a computer without a person checking it. We will not share your information just because it is interesting, and you should not post private information about yourself or others online. Your parent or legal guardian usually helps exercise your privacy rights while you are under 18. You can still ask questions or tell a trusted staff member about a concern.")
    add_p(doc, "Contact: privacy@ruzawi.com or speak to [DPO / trusted staff role]. If you are in immediate danger or need safeguarding help, speak to a trusted adult straight away.")
    add_h(doc, "8. Age-18 / leaver trigger", 1)
    add_p(doc, "On a pupil's 18th birthday or earlier departure, the designated owner shall review data and permissions: delete what is no longer necessary; retain only under an approved legal/operational basis and period; refresh consent directly with the now-adult data subject where future optional processing is desired; record the decision and notify relevant systems/processors.")
    add_sources(doc, [("Children's personal information implementation guideline (Veritas copy of CDPG 2 of 2024)", "https://www.veritaszim.net/sites/veritas_d/files/02%20Processing%20of%20Children%27s%20Personal%20Information.pdf")])
    return save(doc, "04-Consent-and-Childrens-Data-Procedure.docx")


def build_dsr():
    doc = new_doc("Data Subject Rights Handling Procedure", "Request intake, verification, search, decision and closure", "DP-05")
    add_h(doc, "1. Ownership and intake", 1)
    add_p(doc, "The DPO owns this procedure. privacy@ruzawi.com must be monitored, protected by MFA, limited to authorised personnel and covered by a trained alternate. Any employee who receives a request must forward it to the DPO immediately without requiring special wording or a form.")
    add_h(doc, "2. Rights covered", 1)
    add_bullets(doc, [
        "information about use; access; objection to all or part of processing; correction of false, misleading or outdated information; deletion of false or misleading information; consent withdrawal; and objection to direct marketing;",
        "complaints about fairness, security or use; and requests made for a child by a parent/legal guardian or by another authorised representative.",
    ])
    add_h(doc, "3. Procedure and control times", 1)
    add_table(doc, ["Stage","Owner","Control time","Required record"], [
        ("Log and acknowledge","DPO/alternate","Same or next working day","Reference, received date, requester, channel, scope and 30-day target."),
        ("Verify identity/authority","DPO","Promptly; proportionate to risk","Method, evidence sighted, result and any restrictions."),
        ("Clarify scope if genuinely needed","DPO","Without delay","Neutral clarification and revised search terms; do not narrow unfairly."),
        ("Search and preserve","Information owners / IT","Normally within 10 working days","Systems, custodians, date ranges, search results, legal holds."),
        ("Review and decide","DPO + legal/owner as needed","Allow time for redaction and approval","Decision rationale, exemptions/third-party protection, actions."),
        ("Respond and close","DPO","Target within 30 calendar days","Secure response, delivery confirmation, action completion and closure date."),
    ], [1900,1650,1550,4260], font=8.2)
    add_h(doc, "4. Verification and secure response", 1)
    add_bullets(doc, [
        "Use the least intrusive reliable method; do not request identity documents through insecure channels or retain copies unnecessarily.",
        "For a child, verify the adult's parental/legal-guardian authority and consider the child's best interests and any safeguarding restrictions.",
        "Protect information about third parties; redact or obtain consent where appropriate and record the balancing decision.",
        "Use a secure delivery channel appropriate to sensitivity; verify the destination and password exchange separately where encryption is used.",
    ])
    add_h(doc, "5. Search instruction template", 1)
    for label, hint in [
        ("Request reference and due date", "Do not include unnecessary identity documents."),
        ("Data subject / authorised representative", "State verified identifiers needed for a reliable search."),
        ("Scope, date range and search terms", "Include aliases, email addresses, pupil/staff IDs and relevant locations."),
        ("Systems and custodians to search", "Email, SIS, files, cloud, finance, HR, medical/pastoral, CCTV, paper and processors as applicable."),
        ("Return format and deadline", "Preserve metadata and context; flag third-party/safeguarding/legal issues."),
        ("Search completed by / date / no-result explanation", "Record repositories searched and any inaccessible sources."),
    ]:
        form_line(doc, label, hint)
    add_h(doc, "6. Closure quality check", 1)
    add_bullets(doc, [
        "Identity and authority verified; scope and rights correctly characterised.",
        "All relevant owners/systems/processors searched; no legal hold breached.",
        "Decision lawful, consistent and documented; corrections/deletions/objections completed across systems.",
        "Response clear, secure and on time; complaint/POTRAZ route stated where appropriate.",
        "Register updated with evidence location and lessons/remediation.",
    ])
    add_sources(doc)
    return save(doc, "05-Data-Subject-Rights-Handling-Procedure.docx")


def build_breach():
    doc = new_doc("Data Breach Response Plan", "Detection, containment, notification, investigation and recovery", "DP-06")
    add_callout(doc, "Emergency rule", "Immediately report any suspected loss, misdirection, unauthorised access/disclosure/alteration, malware, system compromise or unavailability involving personal information to privacy@ruzawi.com and [IT SECURITY CONTACT]. Do not investigate alone, delete evidence or delay because facts are incomplete.", "danger")
    add_h(doc, "1. Incident command and contacts", 1)
    add_table(doc, ["Role","Primary","Alternate","Responsibility"], [
        ("Incident Lead / DPO","[NAME / MOBILE]","[NAME / MOBILE]","Own timeline, classification, regulator/data-subject notice, records."),
        ("Technical Lead","[NAME / MOBILE]","[NAME / MOBILE]","Contain, preserve evidence, investigate, eradicate and recover."),
        ("Executive Lead","Head: [CONTACT]","Board Chair: [CONTACT]","Decisions, resources, communications, material-risk oversight."),
        ("Information Owner","[DEPARTMENT OWNER]","[ALTERNATE]","Explain data/context; support affected-person analysis and remediation."),
        ("Legal / Communications","[CONTACTS]","[CONTACTS]","Legal privilege/advice; accurate and coordinated messaging."),
    ], [1700,2000,1900,3760], font=8.0)
    add_h(doc, "2. First 24 hours", 1)
    add_table(doc, ["Time","Action","Owner / evidence"], [
        ("0-15 min","Report; record reporter, discovery time, system/data and current impact. Escalate to DPO and IT.","Reporter / incident ticket."),
        ("0-1 hr","Contain safely: revoke sessions, isolate device, stop disclosure, recall email, preserve logs/images. Avoid destroying evidence.","IT / chain-of-custody log."),
        ("0-2 hrs","Open breach register; appoint incident team; determine whether personal information is involved and record awareness time.","DPO / chronology."),
        ("0-6 hrs","Assess data, people, children/sensitive categories, volume, recipients, encryption, likely consequences and ongoing risk.","DPO + owner + IT / risk assessment."),
        ("0-12 hrs","Draft Form DP3 and facts/unknowns; notify insurer/legal/processor as contractually required; plan affected-person notice.","DPO / drafts and correspondence."),
        ("Before 24 hrs","Submit DP3 to POTRAZ for a personal data breach; retain delivery evidence. Do not wait for full investigation.","DPO / submitted form and acknowledgement."),
    ], [1100, 5440, 2820], font=8.0)
    add_h(doc, "3. Risk assessment", 1)
    add_bullets(doc, [
        "Nature, sensitivity and identifiability of data; whether health, biometric, financial, credentials, safeguarding or children's information is involved.",
        "Number and vulnerability of affected people; whether a child or at-risk person could face physical, emotional, identity, financial, disciplinary or reputational harm.",
        "Ease of identification, encryption/protection status, recipient trustworthiness, containment and likelihood of misuse.",
        "Severity, duration, scale and reversibility of harm; ongoing exposure; cross-border or criminal/cyber aspects.",
    ])
    add_h(doc, "4. Notification rules", 1)
    add_table(doc, ["Recipient","Deadline / trigger","Minimum content and evidence"], [
        ("POTRAZ","Within 24 hours of awareness of a personal data breach","Form DP3; facts available, categories/approximate people/records, consequences, measures, DPO contact; document unknowns and updates."),
        ("Affected data subjects","Within 72 hours where likely high risk to rights/freedoms","Clear language: what happened, likely effects, measures taken, practical protective steps, contact and updates. Use accessible/age-appropriate channels."),
        ("Processor/controller partners","Immediately under contract and as necessary","Preserve responsibilities, facts, logs, assistance and subprocessor status."),
        ("zw-CIRT / law enforcement / insurer","As appropriate and authorised","Coordinate without delaying POTRAZ deadline or compromising evidence/data subjects."),
        ("Board / Head","Immediate for material/high-risk events","Decision record, resourcing, communications and residual risk."),
    ], [1900, 3000, 4460], font=8.1)
    add_h(doc, "5. Investigation and recovery", 1)
    add_numbers(doc, [
        "Establish a verified chronology, root cause, systems/accounts affected, exfiltration/access evidence and processor involvement.",
        "Eradicate malicious access, rotate credentials, patch/configure, restore clean data and test availability/security before normal operation.",
        "Respond to POTRAZ information requests within 14 days and maintain a controlled evidence file.",
        "Conclude the investigation and submit a final report within 21 days of initial notification, recording cause, impact, response, notifications, lessons and remediation.",
        "Track corrective actions to closure; update risks, DPIAs, contracts, training, controls and policies; report closure to the Board.",
    ])
    add_h(doc, "6. Incident record / decision form", 1)
    for label in [
        "Incident ID; reporter; discovery and awareness timestamps", "Description, systems, data and affected groups", "Containment actions and evidence preserved", "Risk assessment and high-risk decision", "DP3 submission date/time/reference", "Affected-person notification decision/date/channels", "Root cause and final impact", "21-day final report reference", "Remediation owners/dates and closure approval",
    ]:
        form_line(doc, label)
    add_h(doc, "7. Tabletop exercise scenario", 1)
    add_p(doc, "Scenario: a staff mailbox is compromised and an attacker downloads a spreadsheet containing pupil names, guardian contacts, health flags and fee information. At 08:15 IT sees suspicious forwarding rules; at 09:00 the staff member confirms an unexpected login. The exercise must test awareness-time recording, containment, child/high-risk assessment, DP3 preparation, 72-hour communication, processor cooperation, parent help guidance, recovery and Board reporting.")
    add_p(doc, "Exercise date: ______  Participants: __________________  Actions/owners/dates: __________________  DPO closure: __________________")
    add_sources(doc, [("Form DP3 is contained in the Fourth Schedule to S.I. 155 of 2024", "https://www.potraz.gov.zw/wp-content/uploads/2025/02/sI-155-of-2024-Cyber-and-Data-Protection-Normal_240913_1250178.pdf")])
    return save(doc, "06-Data-Breach-Response-Plan.docx")


def build_security():
    doc = new_doc("Information Security and Access Control Policy", "Minimum technical, organisational and physical safeguards", "DP-07", "IT Security Lead / DPO")
    add_h(doc, "1. Security objectives", 1)
    add_p(doc, "Ruzawi shall protect confidentiality, integrity, availability and recoverability in proportion to the nature of personal information, risks to data subjects, current technology and reasonable implementation cost. Controls apply to school and personal devices used for work, cloud services, email, paper records, CCTV and processors.")
    add_h(doc, "2. Mandatory control baseline", 1)
    add_table(doc, ["Domain","Minimum control","Evidence"], [
        ("Identity","Unique accounts; no sharing; MFA for email, admin, SIS, cloud and remote access where available; prompt joiner/mover/leaver changes.","Account list, MFA report, access tickets, leaver disable times."),
        ("Access","Least privilege; owner approval; termly review of sensitive systems and quarterly review of privileged accounts.","Role matrix, approvals, review sign-off and removals."),
        ("Devices","Managed configuration; encryption; screen lock; patching; endpoint protection; remote wipe where appropriate.","MDM/endpoint reports, encryption and patch compliance."),
        ("Email/web","Anti-phishing/spam controls; safe sharing; verify recipients; disable risky forwarding; secure forms; TLS.","Configuration screenshots, simulated-phishing results, incidents."),
        ("Network/cloud","Secure configuration, restricted administration, logging, vendor review, vulnerability remediation and data-residency record.","Configuration baseline, logs, scan/remediation reports, contracts."),
        ("Backups","Protected backups proportionate to criticality; separation from production; tested restoration.","Backup success reports and restoration-test record."),
        ("Logging","Log authentication, admin and relevant data-access events; protect logs; review alerts; time synchronisation.","Retention/configuration, alert review and investigation tickets."),
        ("Paper/physical","Locked records; controlled keys; clear desk; visitor management; secure transport; fire/water protection where required.","Key register, records-room/visitor logs, inspections."),
        ("CCTV","Purpose/signage; named access; export log; secure credentials; approved retention/overwrite; disclosure controls.","Access list/log, settings, signage, DPIA and deletion checks."),
        ("Disposal","Cross-cut shredding/approved contractor; secure wipe or physical destruction; chain of custody and certificates.","Disposal register, wipe logs, certificates."),
    ], [1500, 5000, 2860], font=7.8)
    add_h(doc, "3. Data classification", 1)
    add_table(doc, ["Class","Examples","Handling"], [
        ("Restricted","Health/safeguarding, credentials, biometrics, identity documents, disciplinary/legal matters","Named access only; strong authentication; encrypted transfer/storage where proportionate; no uncontrolled personal email/device; log access where feasible."),
        ("Confidential","Pupil records, staff/HR, fees, applications, guardian contacts, internal reports","Role-based access; approved systems; secure sharing; retention and disposal controls."),
        ("Internal","Operational material not intended for public release","School accounts/systems; no unauthorised external sharing."),
        ("Public","Approved website, prospectus, published notices","Confirm approval, accuracy, image/consent rights and removal process."),
    ], [1500, 3300, 4560], font=8.1)
    add_h(doc, "4. Access review form", 1)
    for label in ["System / information owner / review date", "Users and roles exported from source system", "Privileged/service/shared accounts reviewed", "Leavers/movers/inactive accounts removed or justified", "MFA and authentication exceptions", "Excess access removed; ticket/reference", "Owner and DPO/IT sign-off"]:
        form_line(doc, label)
    add_h(doc, "5. CCTV retention schedule and operating rules", 1)
    add_table(doc, ["Record / event","Retention","Control and disposal"], [
        ("Routine CCTV footage","30 days from recording","Automatic overwrite/deletion. IT/Security verifies the configured period monthly and records the check."),
        ("Exported footage for a security, safeguarding, disciplinary or safety incident","Until the matter, investigation, claim and any applicable appeal/hold are closed, then promptly delete unless a longer legal/regulator period is documented","Assign incident/case reference; restrict to named personnel; encrypt/protect the export; record every access/disclosure; review at least every 90 days."),
        ("Footage requested by law enforcement, POTRAZ, insurer, court or data subject","Preserve only the relevant segment while the authorised request, legal duty or dispute is active","DPO/legal review before disclosure; preserve original integrity; disclose the minimum necessary; document authority, recipient and date."),
        ("CCTV access and export logs","6 years proposed from entry/closure, subject to legal review","Restricted compliance record; securely delete at expiry."),
    ], [2600, 2700, 4060], font=7.9)
    add_bullets(doc, [
        "CCTV must have a defined safety/security purpose, visible signage, an approved camera map, restricted live/playback access and a completed DPIA/high-risk children's-processing review.",
        "Do not use audio recording, facial recognition, biometric identification or automated behavioural profiling unless separately justified, assessed and approved by the DPO, Head/Board and POTRAZ where required.",
        "An information owner may place a documented hold before routine overwrite. The hold must identify the exact time/camera segment, reason, owner, review date and release trigger; blanket or indefinite preservation is prohibited.",
        "Exports must use approved encrypted storage, a unique case reference and a disclosure log. Personal devices, consumer messaging and uncontrolled removable media are prohibited.",
        "Security and the DPO shall review camera purpose, access list, retention configuration, signage, exports, requests and open holds at least annually and after any incident or system change.",
    ])
    add_h(doc, "6. Document-shredding policy", 1)
    add_p(doc, "Paper containing personal, confidential or restricted information must never be placed in ordinary waste or recycling. It must be placed promptly in a locked confidential-waste console or destroyed by an approved cross-cut shredder so that the information cannot practicably be reconstructed.")
    add_table(doc, ["Stage","Required control","Evidence"], [
        ("Authorise","Confirm the approved retention period has expired and no legal, safeguarding, complaint, investigation, incident or regulator hold applies.","Batch approval and hold check."),
        ("Collect","Use locked, clearly labelled confidential-waste consoles in controlled locations. Do not leave bags, boxes or loose records unattended.","Console/location register and collection log."),
        ("Destroy internally","Use an approved cross-cut shredder; authorised staff supervise; re-shred any readable fragments; dispose of residue securely.","Batch ID, operator, witness, date and approximate volume."),
        ("Destroy by contractor","Use an approved, contracted and vetted provider; maintain sealed chain of custody; require secure transport and witnessed/verified destruction.","Collection receipt, chain-of-custody log and certificate of destruction."),
        ("Close","Record category/date range, retention trigger, volume, method, exceptions and completion; investigate any loss, spillage or failed destruction as a possible breach.","Disposal register and DPO/records-owner closure."),
    ], [1500, 5200, 2660], font=8.0)
    add_p(doc, "Small quantities may be shredded by authorised staff immediately. Bulk archives and highly sensitive records should use the approved contractor unless the Records Owner and DPO approve an equally secure internal method. The School must sample certificates, contractor performance and confidential-waste locations at least annually.")
    add_h(doc, "7. Security testing and exceptions", 1)
    add_p(doc, "IT must test control effectiveness and document improvement: backup restoration, incident/tabletop exercises, patch/vulnerability review, access review, phishing awareness and selected configuration audits. Any exception requires scope, reason, risk, compensating control, owner, approval and expiry date; indefinite exceptions are prohibited.")
    add_h(doc, "8. Staff rules", 1)
    add_bullets(doc, [
        "Lock screens, protect credentials, use approved storage and verify recipients/links/requests.",
        "Do not download or copy personal information to personal devices, consumer cloud storage or messaging services without written approval.",
        "Do not photograph, print or discuss pupil/staff information unnecessarily; collect printouts promptly and dispose securely.",
        "Report mistakes, suspicious messages, lost devices, misdirected email and unusual access immediately.",
    ])
    add_sources(doc)
    return save(doc, "07-Information-Security-and-Access-Control-Policy.docx")


def build_retention():
    doc = new_doc("Data Retention and Secure Disposal Schedule", "Category-specific draft periods, triggers, legal holds and destruction evidence", "DP-08")
    add_callout(doc, "Validation required", "The periods below are defensible operational starting points, not a substitute for Zimbabwean legal review. Information owners must identify applicable education, employment, tax, safeguarding, health, limitation and archival rules before approval.", "warning")
    add_h(doc, "1. Rules", 1)
    add_bullets(doc, [
        "Retain identifiable information only as long as necessary for the recorded purpose and law; apply the shortest applicable period.",
        "A legal hold, active complaint, investigation, safeguarding need or litigation may suspend disposal; record scope, approver and release date.",
        "At age 18 or pupil departure, review data given under guardian consent and delete, retain under another basis or refresh consent directly with the adult.",
        "Disposal must cover live systems, shared drives, email, paper, devices and processors; backup expiry may follow the approved backup cycle if access is restricted and restoration controls prevent ordinary use.",
    ])
    add_h(doc, "2. Draft retention schedule", 1)
    rows = [
        ("Admissions - unsuccessful/withdrawn","Close of process or waiting-list expiry","12 months, unless dispute/consent/other documented need","Delete files, database entry and working copies; retain minimal outcome/audit record if justified."),
        ("Admissions - enrolled","Enrolment","Transfer to pupil record; apply pupil categories","Remove duplicate admissions copies after verification."),
        ("Pupil core academic record","Departure","Define permanent/minimum archive after education/legal review","Archive/minimise; separate optional publicity consent."),
        ("Health/sanatorium and safeguarding","Last interaction / departure","Case-specific statutory/professional/safeguarding period; legal review mandatory","Restricted archive; secure destruction after hold/period."),
        ("CCTV","Recording","Normally 30 days unless approved shorter/longer need or incident hold","Automatic overwrite; export only for incident; log access/export/deletion."),
        ("Visitor/access logs","Visit / event","12 months unless security incident/legal need","Secure deletion or shredding."),
        ("Staff recruitment - unsuccessful","Decision","12 months","Delete CV, notes and references; retain minimal equality/legal record if applicable."),
        ("Staff personnel/payroll","Termination / transaction","6 years proposed minimum for core employment/financial records; validate law","Secure archive then shred/wipe; retain only required references/benefits records."),
        ("Finance/tax/audit","Financial year end","6 years proposed; validate tax/accounting rules","Secure destruction after audit/legal hold."),
        ("Rights requests/complaints","Closure","6 years proposed for accountability/claims; minimise ID evidence","Delete working copies; retain decision, correspondence and proof of action."),
        ("Breach/incident records","Closure/final report","6 years proposed; longer for material/regulator-directed cases","Restricted incident archive; preserve notification evidence."),
        ("Website contact enquiries","Closure","12 months unless converted to relationship","Delete inbox/system copies and attachments."),
        ("Consent records","Withdrawal/expiry/end of related processing","Duration of processing plus 6 years proposed for proof","Retain wording/version/timestamp and withdrawal outcome, not unnecessary content."),
        ("Training records","Completion","Employment/engagement plus 3 years proposed","Retain completion evidence; delete unnecessary training analytics."),
        ("Vendor/DPA/DPIA/risk/audit","End of relationship/control","Contract life plus 6 years proposed; superseded DPIAs retained for audit","Restricted governance archive."),
        ("Analytics/cookie data","Collection/choice","Exact vendor/account and consent-tool settings; minimise","Configure expiry/deletion; respect withdrawal for future collection."),
        ("Security logs/backups","Event/backup","Risk-based: commonly 30-180 days for logs; backup cycle per recovery need","Automated expiry; restricted access; tested deletion/rotation."),
    ]
    add_table(doc, ["Category","Trigger","Draft period","Disposal / control"], rows, [2300, 1700, 2300, 3060], font=7.3)
    add_h(doc, "3. Disposal procedure", 1)
    add_numbers(doc, [
        "Owner identifies records due and confirms no legal, safeguarding, complaint, incident or regulator hold.",
        "DPO/records owner approves batch and method; processor deletion is requested where applicable.",
        "Paper is cross-cut shredded or destroyed by approved contractor; media is securely wiped using an approved method or physically destroyed when wiping is unreliable.",
        "System data is removed from live and ordinary-access copies; backup expiry is documented; access to residual protected backups remains prohibited except controlled restoration.",
        "Record date, category, period/trigger, volume, systems/media, method, operator, witness/contractor certificate and exceptions in the disposal register.",
    ])
    add_h(doc, "4. Disposal record", 1)
    for label in ["Batch ID / owner / approval date", "Category, date range and volume", "Retention trigger and hold check", "Systems, paper boxes, devices or processors", "Method and completion date", "Operator / witness / contractor certificate", "Exceptions or failed deletions and remediation", "DPO closure"]:
        form_line(doc, label)
    add_sources(doc)
    return save(doc, "08-Data-Retention-and-Secure-Disposal-Schedule.docx")


def build_vendor_transfer():
    doc = new_doc("Processor, DPA and Cross-Border Transfer Pack", "Due diligence, contracting, transfer review and minimum clauses", "DP-09", "DPO / Procurement / IT")
    add_h(doc, "1. Pre-contract gate", 1)
    add_numbers(doc, [
        "Define purpose, necessity, data, children/sensitive categories, users, access, countries, retention, security and exit needs.",
        "Complete DPIA screening and children's high-risk classification before procurement or trial data is shared.",
        "Perform due diligence on the processor and all material subprocessors; validate hosting/support/backup locations.",
        "Complete transfer review and obtain required POTRAZ notification/authorisation before any cross-border processing begins.",
        "Execute DPA and transfer terms; approve risk and residual exceptions; record contract owner and renewal/exit dates.",
        "Test configuration, minimise fields/access, train users and retain sign-off before go-live.",
    ])
    add_h(doc, "2. Vendor due-diligence checklist", 1)
    add_table(doc, ["Area","Required response/evidence","Result"], [
        ("Entity and service","Legal name, role, service/purpose, data flow, contract entity, support model","[Pass / gap]"),
        ("Children and sensitive data","Specific protections, age-appropriate design, default settings, staff vetting, prohibition on own-purpose use","[Pass / gap]"),
        ("Locations/transfers","Primary/backup/support countries, remote access, subprocessors, transfer mechanism, government-access approach","[Pass / gap]"),
        ("Security","Certifications, encryption, MFA, least privilege, logging, vulnerability/patching, testing, continuity and deletion","[Pass / gap]"),
        ("Incidents","Detection, immediate notice, 24-hour regulator support, evidence preservation, investigation and communication assistance","[Pass / gap]"),
        ("Rights/records","Search/export/correct/delete capability; audit logs; assistance; data return and deletion at exit","[Pass / gap]"),
        ("Subprocessors","Current list, notice/objection, equivalent obligations and liability","[Pass / gap]"),
        ("Commercial/legal","Insurance, audit rights, liability, termination, service changes, dispute/enforcement","[Pass / gap]"),
    ], [1800, 6200, 1360], font=8.0)
    add_h(doc, "3. Minimum DPA clauses", 1)
    add_p(doc, "The contract must be adapted to the transaction and reviewed by counsel. At minimum it should state:")
    add_bullets(doc, [
        "subject matter, duration, nature, purpose, data types, data subjects and documented controller instructions;",
        "confidentiality and authorised personnel; appropriate technical, organisational and physical security; evidence and testing;",
        "no processing for the processor's own advertising, profiling, model training or unrelated purpose without express lawful written agreement;",
        "subprocessor approval/notice, equivalent written terms and continuing processor responsibility;",
        "international locations, transfer mechanism, POTRAZ authorisation conditions and no location change without written approval;",
        "prompt incident notice with all available facts and ongoing cooperation sufficient for Ruzawi's 24/72-hour and 21-day duties;",
        "assistance with rights, DPIAs, regulator inquiries, audits, records, corrections, deletion and restriction;",
        "return/deletion at termination, including backups according to a controlled cycle, plus written certification and transition support;",
        "audit/information rights, remediation, liability/indemnity, insurance, termination for material privacy/security breach and governing law/disputes.",
    ])
    add_h(doc, "4. Cross-border transfer review", 1)
    for label, hint in [
        ("Transfer ID, purpose and owner", "Link to ROPA, DPIA and vendor contract."),
        ("Data subjects/data and children's/high-risk classification", "Identify sensitive, health, biometric, safeguarding and credentials."),
        ("Exporter/importer and all countries", "Include storage, backup, support, remote access and subprocessors."),
        ("Legal condition/safeguard", "Adequacy, unambiguous consent, contract necessity, SCC/IDTA/equivalent or other condition; attach evidence."),
        ("Recipient-country and vendor risk assessment", "Law, access, enforceability, security, onward transfer and remedies."),
        ("POTRAZ authorisation/notification", "Reference, date, scope, conditions or documented exemption. Children's transfers require enhanced prior review."),
        ("Supplementary controls", "Encryption/key control, pseudonymisation, minimisation, access, logs, deletion."),
        ("Approval, review trigger and expiry", "DPO, owner, Head/Board where required."),
    ]:
        form_line(doc, label, hint)
    add_h(doc, "5. Initial website-provider register - validation required", 1)
    add_table(doc, ["Provider","Purpose","Key evidence to obtain"], [
        ("Netlify","Hosting and serverless form handling","DPA; hosting/log locations; security; subprocessors; deletion; incident terms; transfer mechanism."),
        ("Resend","Transactional email","DPA; message/log retention; countries; subprocessors; security; deletion/incident support."),
        ("Google reCAPTCHA","Form abuse prevention","Terms/privacy; data flow; countries; necessity configuration; children's/high-risk review; transfer approval."),
        ("CookieHub","Consent management","DPA/terms; consent record data/retention; countries; scan/config evidence."),
        ("Google Analytics","Consented analytics","DPA/terms; property retention; IP/signals; countries; children/account settings; transfer review."),
        ("Microsoft Clarity","Consented analytics/session replay","DPA/terms; masking/exclusion verification; retention; countries; transfer review."),
        ("Application database provider","Admissions data and uploads","Highest priority: DPA, hosting/support/backup countries, child/sensitive controls, DPIA, deletion and POTRAZ authorisation."),
    ], [2000, 2500, 4860], font=8.0)
    add_sources(doc, compact_review=True)
    return save(doc, "09-Processor-DPA-and-Cross-Border-Transfer-Pack.docx")


def build_dpia_risk():
    doc = new_doc("DPIA and Privacy Risk Management Pack", "Screening, assessment, risk scoring, approval and review", "DP-10")
    add_h(doc, "1. Mandatory screening triggers", 1)
    add_bullets(doc, [
        "children's or sensitive/health/biometric/genetic information; CCTV, surveillance, tracking or systematic monitoring;",
        "international transfer, cloud hosting or remote support involving children; new/innovative technology or large-scale processing;",
        "profiling or automated decisions affecting placement, discipline, access or another significant outcome;",
        "data matching, new purpose, new processor, new sharing, major field expansion or a change likely to increase harm.",
    ])
    add_h(doc, "2. DPIA process", 1)
    add_numbers(doc, [
        "Describe the proposal, purpose, necessity, data flow, systems, people, recipients, countries, retention and legal basis.",
        "Consult relevant stakeholders, including child/safeguarding perspectives and processor/security specialists where appropriate.",
        "Assess necessity, proportionality, transparency, consent/rights, minimisation, best interests and alternatives.",
        "Identify threats and harms to people; score inherent likelihood and impact; define controls and residual risk.",
        "Determine POTRAZ notification/authorisation and transfer requirements; do not go live until completed.",
        "DPO records advice; accountable owner accepts only permissible residual risk; high unresolved risk is escalated to Head/Board and POTRAZ as required.",
        "Review annually and on change, incident, complaint, vendor/location change, new evidence or control failure.",
    ])
    add_h(doc, "3. Risk scoring", 1)
    add_table(doc, ["Score","Likelihood","Impact"], [
        ("1","Rare / strong prevention and no known occurrence","Negligible inconvenience; no meaningful rights effect."),
        ("2","Unlikely","Limited, short-lived or easily reversible harm."),
        ("3","Possible","Material distress, confidentiality loss, service/rights impact or moderate financial/reputational harm."),
        ("4","Likely","Serious or prolonged harm; sensitive/child data; significant identity, safety, financial or discrimination risk."),
        ("5","Almost certain","Severe, widespread or irreversible harm; safeguarding danger or systemic compromise."),
    ], [900, 3900, 4560], font=8.2)
    add_p(doc, "Risk score = likelihood x impact. Suggested bands: 1-4 low; 5-9 medium; 10-15 high; 16-25 critical. The DPO may require stronger treatment based on the nature of the affected people/data even where arithmetic is lower.")
    add_h(doc, "4. DPIA template", 1)
    for label, hint in [
        ("Project, owner, DPO and dates", "Include planned go-live and review trigger."),
        ("Purpose, expected benefit and necessity", "Explain why a less intrusive option cannot achieve the purpose."),
        ("Data flow", "Collection -> use -> storage -> sharing/transfer -> archive/deletion; attach diagram."),
        ("Data subjects and categories", "Flag children, sensitive/health/biometric data and vulnerable groups."),
        ("Legal basis, notices, consent and rights", "Explain guardian verification, withdrawal and human review where relevant."),
        ("Systems, processors, subprocessors and countries", "Link contracts, due diligence, authorisation and transfer review."),
        ("Retention, minimisation and access", "Fields, optional/required status, roles, logging and deletion."),
        ("Consultation", "Pupils/parents/staff/safeguarding/IT/vendor/legal as proportionate."),
        ("Risks and mitigations", "Use risk register: cause, event, harm, inherent score, controls, residual score, owner/date."),
        ("POTRAZ engagement", "Notification/authorisation reference, conditions or reason not required."),
        ("Decision", "Approve / approve with actions / do not proceed; DPO advice; owner acceptance; Head/Board/POTRAZ escalation."),
    ]:
        form_line(doc, label, hint)
    add_h(doc, "5. Initial privacy risk register", 1)
    add_table(doc, ["ID","Risk","Inherent","Key treatment / evidence","Owner / residual"], [
        ("R-01","Children's admissions data is transferred to cloud/email/database providers without complete authorisation or safeguards","5x5=25","Map countries/subprocessors; DPIA; DPAs; POTRAZ authorisation; minimise email content; encryption/access/deletion evidence","DPO/IT/Admissions - rescore"),
        ("R-02","Guardian self-declaration is accepted without reliable authority verification","4x5=20","Sight qualifying evidence; record result/restrictions; safeguarding escalation; minimise copies","Admissions/DPO - rescore"),
        ("R-03","Clarity/session replay captures form or sensitive content despite intended masking","3x5=15","Test production masking/exclusions; disable on form routes; consent test; vendor evidence and monitoring","Web/IT/DPO - rescore"),
        ("R-04","Compromised staff email exposes pupil/guardian information","4x5=20","MFA; forwarding alerts; least privilege; phishing training; incident plan; logging and access reviews","IT - rescore"),
        ("R-05","CCTV access or retention exceeds purpose","3x4=12","DPIA; named access; export log; 30-day overwrite or approved period; signage; authorisation review","Security/DPO - validate"),
        ("R-06","No consistent deletion across systems/processors/backups","4x4=16","Approve schedule; configure lifecycle; processor deletion; disposal logs; sampling/audit","Owners/IT - rescore"),
        ("R-07","Rights request misses systems or deadline","3x4=12","Central log, trained alternates, system map, due-date alerts, closure QC and metrics","DPO - rescore"),
        ("R-08","Breach discovered but 24-hour DP3 deadline is missed","3x5=15","24/7 escalation, awareness definition, prefilled DP3, tabletop exercise, executive contacts","DPO/IT - rescore"),
    ], [650, 2600, 1000, 3740, 1370], font=7.4)
    add_sources(doc)
    return save(doc, "10-DPIA-and-Privacy-Risk-Management-Pack.docx")


def build_training():
    doc = new_doc("Training, Awareness and Induction Pack", "Role-based curriculum, completion evidence and refresher controls", "DP-11", "DPO / HR")
    add_h(doc, "1. Training standard", 1)
    add_p(doc, "All employees, governors, volunteers and regular contractors with access to personal information must complete induction before or promptly after access is granted and annual refresher training. DPO, IT, admissions, HR, finance, medical/pastoral, safeguarding, CCTV/security and senior leaders require role-specific training.")
    add_h(doc, "2. Core curriculum", 1)
    add_table(doc, ["Module","Audience","Learning outcome","Evidence"], [
        ("Privacy foundations","All personnel","Recognise personal/sensitive/children's data; apply principles, purpose and minimisation.","Attendance/quiz/acknowledgement."),
        ("Children and consent","Admissions, teaching, pastoral, medical, activities","Guardian authority, age-appropriate notices, consent/withdrawal, best interests and age-18 review.","Scenario assessment and sign-off."),
        ("Rights requests","All + DPO/owners","Recognise any-form request; forward immediately; verify/search/respond securely and on time.","Exercise and register simulation."),
        ("Security and safe handling","All","MFA/passwords, phishing, approved systems, sharing, paper, devices and clean desk.","Quiz/phishing metric."),
        ("Breach response","All + incident team","Report immediately; contain safely; preserve evidence; understand 24/72-hour and 21-day duties.","Tabletop/attendance/actions."),
        ("DPIA and procurement","Leaders, IT, procurement, owners","Screen before change; processors/transfers; authorisation; no go-live without approval.","Completed sample screening."),
        ("CCTV/physical records","Security, admin, records staff","Access/export logs, signage, retention, disclosure and destruction.","Practical inspection/sign-off."),
        ("DPO/Board governance","DPO, Head, Board","Registration, oversight, audits, risks, regulator cooperation and evidence.","Briefing minute / certification."),
    ], [1700, 1700, 4200, 1760], font=7.8)
    add_h(doc, "3. New-starter induction checklist", 1)
    add_bullets(doc, [
        "[ ] Privacy Policy and public notices explained; confidentiality obligations acknowledged.",
        "[ ] Role-specific systems and least-privilege access approved; MFA enabled; prohibited tools explained.",
        "[ ] Children/sensitive-data, safeguarding and consent rules understood.",
        "[ ] Rights requests must be forwarded immediately to privacy@ruzawi.com.",
        "[ ] Incidents, misdirected email, lost devices and suspicious access must be reported immediately.",
        "[ ] Retention, printing, records-room, CCTV and secure disposal rules relevant to role explained.",
        "[ ] Knowledge check completed; questions resolved; acknowledgement signed.",
    ])
    form_line(doc, "Starter / role / start date")
    form_line(doc, "Trainer / completion date / score")
    form_line(doc, "Starter acknowledgement and signature")
    add_h(doc, "4. Annual training plan", 1)
    add_table(doc, ["Quarter","Activity","Target","Owner / evidence"], [
        ("Q1","Annual all-staff refresher and policy acknowledgement","100% active personnel","DPO/HR; attendance, quiz and overdue escalation."),
        ("Q2","Rights-request and admissions/guardian scenario workshop","DPO, Admissions, HR, Finance, Pastoral","Exercise outputs and actions."),
        ("Q3","Phishing/security awareness plus CCTV/records inspection","All; IT/security/admin focus","Campaign results, inspection and remediation."),
        ("Q4","Breach tabletop and Board/DPO annual compliance briefing","Incident team, leadership, Board","Exercise report, minutes and next-year plan."),
    ], [1000, 3600, 2400, 2360], font=8.1)
    add_h(doc, "5. Training register fields", 1)
    add_table(doc, ["Person / role","Module/version","Date / method","Result","Certificate/evidence","Refresher due"], [
        ("[NAME]","[MODULE]","[DATE / IN PERSON or ONLINE]","[PASS / ACTION]","[LINK / FILE REF]","[DATE]"),
        ("[NAME]","[MODULE]","[DATE / IN PERSON or ONLINE]","[PASS / ACTION]","[LINK / FILE REF]","[DATE]"),
    ], [1700, 1750, 1700, 1100, 1850, 1260], font=8.0)
    add_sources(doc)
    return save(doc, "11-Training-Awareness-and-Induction-Pack.docx")


def build_registers():
    doc = new_doc("Operational Registers and Forms", "Controlled templates for ongoing compliance evidence", "DP-12")
    add_callout(doc, "Use", "Maintain these registers in a secure, access-controlled system. Do not store unnecessary sensitive details in summary registers; link to restricted case files by reference. Each owner must review entries at least quarterly.")
    registers = [
        ("1. Consent register", ["Consent ID","Person / child","Guardian authority verified","Purpose / wording version","Given / method","Withdrawal / outcome","Evidence ref"], [900,1300,1500,1800,1200,1400,1260]),
        ("2. Data subject rights register", ["Request ID","Received / due","Requester / authority","Right and scope","Owners/systems","Decision / sent","Closed / evidence"], [900,1300,1500,1700,1500,1300,1160]),
        ("3. Breach register", ["Incident ID","Discovery / awareness","Data / people","Risk","DP3 / 24h","Data subjects / 72h","Final report / actions"], [900,1500,1700,900,1300,1400,1660]),
        ("4. International transfer register", ["Transfer ID","Purpose / ROPA","Exporter / importer","Countries","Data / children","Safeguard / POTRAZ ref","Review / owner"], [900,1400,1500,1000,1400,1900,1260]),
        ("5. DPIA register", ["DPIA ID","Project / owner","Trigger","Inherent / residual risk","POTRAZ ref","Decision / actions","Review due"], [900,1600,1100,1800,1200,1700,1060]),
        ("6. Privacy risk register", ["Risk ID","Cause / event / harm","Inherent","Controls","Residual","Owner / due","Status / evidence"], [800,2200,1000,1800,1000,1300,1260]),
        ("7. Training register", ["Person / role","Module / version","Completion date","Result","Evidence","Refresher due"], [1800,1800,1400,1100,1900,1360]),
        ("8. Disposal register", ["Batch ID","Category / range","Trigger / hold check","System/media","Method/date","Operator/certificate","DPO closure"], [850,1600,1700,1300,1300,1600,1010]),
        ("9. Processor and contract register", ["Vendor / owner","Service / data","DPA date","Countries / transfer ref","Subprocessors","Renewal / exit","Risk / evidence"], [1500,1700,1100,1750,1300,1100,910]),
        ("10. Policy/version control register", ["Document ID","Title","Version","Owner","Approved by/date","Next review","Change summary"], [1100,1800,800,1200,1700,1100,1660]),
    ]
    for heading, headers, widths in registers:
        add_h(doc, heading, 1)
        add_table(doc, headers, [["[ENTRY]" for _ in headers],["[ENTRY]" for _ in headers]], widths, font=7.2)
    add_h(doc, "11. Legitimate-interest / non-consent decision record", 1)
    for label, hint in [
        ("Purpose, data subjects and data", "State whether children or sensitive information are involved."),
        ("Proposed basis and necessity", "Identify exact law/interest/contract/vital/public basis; explain why consent is not used."),
        ("Benefits and reasonable expectations", "Consider relationship, context and transparency."),
        ("Risks and balancing", "Impact, vulnerability, safeguards, opt-out/rights and less intrusive alternatives."),
        ("Children's best interests and POTRAZ notification", "Record legal/DPO advice and notification reference where applicable."),
        ("Decision, limits, owner and review date", "Approve / reject / conditions; link ROPA/DPIA."),
    ]:
        form_line(doc, label, hint)
    add_h(doc, "12. Compliance evidence index", 1)
    add_table(doc, ["Checklist item","Evidence title","File/location","Owner","Date/current version","Reviewer / result"], [
        ("[1-36]","[DOCUMENT / SCREENSHOT / LOG / MINUTE]","[CONTROLLED LOCATION]","[NAME]","[DATE / VERSION]","[DPO / PASS-GAP]"),
        ("[1-36]","[DOCUMENT / SCREENSHOT / LOG / MINUTE]","[CONTROLLED LOCATION]","[NAME]","[DATE / VERSION]","[DPO / PASS-GAP]"),
    ], [1000, 2100, 2000, 1200, 1600, 1460], font=7.8)
    add_sources(doc)
    return save(doc, "12-Operational-Registers-and-Forms.docx")


def main():
    builders = [
        build_status_report,
        build_policy,
        build_governance,
        build_ropa,
        build_consent_children,
        build_dsr,
        build_breach,
        build_security,
        build_retention,
        build_vendor_transfer,
        build_dpia_risk,
        build_training,
        build_registers,
    ]
    paths = [fn() for fn in builders]
    print("\n".join(str(p) for p in paths))


if __name__ == "__main__":
    main()
